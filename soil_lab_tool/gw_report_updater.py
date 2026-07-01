"""
Groundwater Report Updater
==========================
Merges Bactochem PDF lab results into an existing Word monitoring report,
updates the historical data table, regenerates trend charts, and updates
the Mann-Kendall XLS file.

Public API:
    run_update_bytes(word_bytes, lab_pdf_bytes, mk_xls_bytes, field_pdf_bytes=None)
        -> (updated_word_bytes, updated_mk_xls_bytes)
"""

import copy
import re
import tempfile
from datetime import datetime
from io import BytesIO
from pathlib import Path

import numpy as np
import xlrd
import xlwt
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
import pdfplumber
from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor

# ──────────────────────────────────────────────────────────────────────────────
# BACTOCHEM PDF PARSER
# ──────────────────────────────────────────────────────────────────────────────

CAS_TO_HEB = {
    "71-43-2":   "בנזן",
    "100-41-4":  "אתיל בנזן",
    "1634-04-4": "MTBE",
    "108-88-3":  "טולואן",
    "1330-20-7": "קסילן",
}


def parse_aminolab_pdf(pdf_path: str) -> dict:
    """
    Parse a single Aminolab groundwater monitoring certificate PDF.

    NOTE: unlike Bactochem, one Aminolab PDF certificate covers a SINGLE
    well/sample. Use parse_aminolab_pdfs() to merge several certificates
    (one per well) from one sampling round.

    Returns the same shape as parse_bactochem_pdf():
        {
          "sampling_date": "09.06.26",
          "samples": {
            "מת-1": {
              "date": "09.06.26",
              "results": {"בנזן": None, "MTBE": None, "טולואן": 0.035, ...},
              "field": {"pH": 6.6, "EC": 1388.0, "עומק פני המים": 54.25, ...},
              "water_level": 54.25,
              "total_depth": 57.15,
              "floating_layer": False,
            },
          }
        }
    """
    well_re = re.compile(r'^(\d+)\s*-\s*םוהת\s*ימ\s*:\s*המגודה\s*רואת$')
    date_re = re.compile(r'^(\d{2}/\d{2}/\d{4})\s*:םוגידה ךיראת$')
    ph_re   = re.compile(r'^-\s*([\d.]+)\s*-\s*pH הבגה$')
    ec_re   = re.compile(r'^-\s*([\d,]+)\s*[¥µu]S/cm תוכילומ$')
    wl_re   = re.compile(r'^-\s*([\d.]+)\s*M\s+םימה ינפ קמוע\s*$')
    td_re   = re.compile(r'^-\s*([\d.]+)\s*M\s+חודיק קמוע\s*$')
    floating_re = re.compile(r'הפצ הבכש')

    # Lines like: "* Toluene 108-88-3 35 84" / "Xylene's - 40 -" / "Benzene 71 -43 -2 <10 -"
    compound_re = re.compile(
        r'^\*?\s*(?P<name>.+?)\s*\*?\s+(?P<cas>[\d][\d\s\-]*\d|-)\s+'
        r'(?P<result><?\d+(?:\.\d+)?)\s+(?P<qual>-|\d+(?:\.\d+)?)\s*$'
    )

    well_id = None
    sample_date = None
    field = {}
    water_level = None
    total_depth = None
    floating_layer = False
    results = {}

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            for raw_line in text.splitlines():
                line = raw_line.strip()

                if well_id is None:
                    m = well_re.match(line)
                    if m:
                        well_id = f"מת-{m.group(1)}"
                        continue

                if sample_date is None:
                    m = date_re.match(line)
                    if m:
                        try:
                            dt = datetime.strptime(m.group(1), "%d/%m/%Y")
                            sample_date = dt.strftime("%d.%m.%y")
                        except ValueError:
                            pass
                        continue

                m = ph_re.match(line)
                if m:
                    field["pH"] = float(m.group(1))
                    continue

                m = ec_re.match(line)
                if m:
                    field["EC"] = float(m.group(1).replace(",", ""))
                    continue

                m = wl_re.match(line)
                if m:
                    water_level = float(m.group(1))
                    field["עומק פני המים"] = water_level
                    continue

                m = td_re.match(line)
                if m:
                    total_depth = float(m.group(1))
                    continue

                if floating_re.search(line):
                    floating_layer = True
                    continue

                # VOC compound rows (BTEX) — only on pages with "Compound"/"CAS No."
                m = compound_re.match(line)
                if m:
                    name = m.group("name").strip()
                    cas = re.sub(r'\s*-\s*', '-', m.group("cas").strip())
                    heb = CAS_TO_HEB.get(cas)
                    if not heb and "xylen" in name.lower():
                        heb = "קסילן"
                    if heb:
                        raw_result = m.group("result")
                        val = None if raw_result.startswith("<") else float(raw_result)
                        if val is not None:
                            val = val / 1000.0  # ppb -> mg/L
                        results[heb] = val

    if well_id is None:
        well_id = "מת-?"

    sample = {
        "date": sample_date,
        "results": results,
        "field": field,
        "water_level": water_level,
        "total_depth": total_depth,
        "floating_layer": floating_layer,
    }

    return {
        "sampling_date": sample_date,
        "samples": {well_id: sample},
    }


def parse_aminolab_pdfs(pdf_paths: list) -> dict:
    """
    Parse and merge several Aminolab certificates (one per well) from the
    same sampling round into the combined {"sampling_date", "samples"} shape.
    """
    combined_samples = {}
    sampling_date = None

    for pdf_path in pdf_paths:
        lab = parse_aminolab_pdf(pdf_path)
        if sampling_date is None:
            sampling_date = lab["sampling_date"]
        combined_samples.update(lab["samples"])

    return {"sampling_date": sampling_date, "samples": combined_samples}


def parse_bactochem_pdf(pdf_path: str) -> dict:
    """
    Parse a Bactochem groundwater monitoring PDF.

    Returns:
        {
          "sampling_date": "03.08.25",
          "samples": {
            "מת-1": {
              "date": "03.08.25",
              "results": {"בנזן": 0.014, "MTBE": 0.9, ...},  # None = ND
              "field": {"pH": 7.11, "EC": 6640, "עומק פני המים": 15.92, ...},
              "water_level": 15.92,
              "total_depth": 21.62,
              "floating_layer": False,
            },
          }
        }
    """
    samples = {}
    global_date = None

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            lines = text.splitlines()

            # Find all sample-description lines in this page
            # Pattern (RTL-mirrored): "1933818 :המגודה רפסמ 5-תמ הדוהי רוא לונוס-םוהת ימ םוגיד"
            segments = []
            for i, line in enumerate(lines):
                m = re.search(r'(\d+)-תמ', line)
                if m and ("הדוהי" in line or "לונוס" in line or "דלק" in line or "תמ" in line):
                    if "בלנק" not in line and "קיפ" not in line:
                        segments.append((f"מת-{m.group(1)}", i))
            segments.append((None, len(lines)))  # sentinel

            for seg_idx in range(len(segments) - 1):
                well_id, start = segments[seg_idx]
                _, end = segments[seg_idx + 1]
                seg_lines = lines[start:end]

                if well_id in samples:
                    continue  # already parsed (blank sample continuation)

                sample = {
                    "date": None,
                    "results": {},
                    "field": {},
                    "water_level": None,
                    "total_depth": None,
                    "floating_layer": False,
                }

                for line in seg_lines:
                    # Sampling date: "03/08/2025 07:20 :םוגיד דעומ"
                    m = re.search(r'(\d{2}/\d{2}/\d{4})\s+\d{2}:\d{2}', line)
                    if m and sample["date"] is None:
                        try:
                            dt = datetime.strptime(m.group(1), "%d/%m/%Y")
                            sample["date"] = dt.strftime("%d.%m.%y")
                            if global_date is None:
                                global_date = sample["date"]
                        except ValueError:
                            pass

                    # Analyte: "CAS #: 71-43-2 0.001 mg/L 0.720 MTBE"
                    m = re.search(
                        r'CAS #:\s*([\d\-]+)\s+[\d.]+\s+mg/L\s+([\d.]+|Not Detected)',
                        line, re.IGNORECASE
                    )
                    if m:
                        cas, raw_val = m.group(1), m.group(2)
                        heb = CAS_TO_HEB.get(cas)
                        if heb:
                            sample["results"][heb] = (
                                None if raw_val.lower() == "not detected"
                                else float(raw_val)
                            )

                    # Floating layer: "ןכ : הפצ הבכש תוחכונ"
                    if "הפצ הבכש" in line and "ןכ" in line:
                        sample["floating_layer"] = True

                    # Field: water level "M 15.59 ןוילע סלפמ קמוע"
                    m = re.search(r'M\s+([\d.]+)\s+ןוילע סלפמ קמוע', line)
                    if m:
                        v = float(m.group(1))
                        sample["water_level"] = v
                        sample["field"]["עומק פני המים"] = v

                    # EC: "µS/cm 6640 רחאל) תילמשח תוכילומ"
                    m = re.search(r'µS/cm\s+([\d,]+)', line)
                    if m:
                        sample["field"]["EC"] = float(m.group(1).replace(",", ""))

                    # pH: "pH units 7.11"
                    m = re.search(r'pH units\s+([\d.]+)', line)
                    if m:
                        sample["field"]["pH"] = float(m.group(1))

                    # DO: "mg/L 2.1 (תובצייתה רחאל) סמומ ןצמח"
                    m = re.search(r'mg/L\s+([\d.]+)\s+\(תובצייתה רחאל\) סמומ ןצמח', line)
                    if m:
                        sample["field"]["חמצן מומס"] = float(m.group(1))

                    # Redox: "mV -229 (תובצייתה רחאל) סקודר"
                    m = re.search(r'mV\s+(-?[\d.]+)\s+\(תובצייתה רחאל\) סקודר', line)
                    if m:
                        sample["field"]["רדוקס"] = float(m.group(1))

                    # Temp: "C° 27.9 (תובצייתה רחאל) הרוטרפמט"
                    m = re.search(r'C°\s+([\d.]+)\s+\(תובצייתה רחאל\) הרוטרפמט', line)
                    if m:
                        sample["field"]["טמפרטורה"] = float(m.group(1))

                    # Turbidity: "NTU 18 (תובצייתה רחאל) תוריכע"
                    m = re.search(r'NTU\s+([\d.]+)\s+\(תובצייתה רחאל\) תוריכע', line)
                    if m:
                        sample["field"]["עכירות"] = float(m.group(1))

                    # Total depth: "M 21.62 חודיקה לש יללכ קמוע"
                    m = re.search(r'M\s+([\d.]+)\s+חודיקה לש יללכ קמוע', line)
                    if m:
                        v = float(m.group(1))
                        sample["total_depth"] = v
                        sample["field"]["עומק כללי של הקידוח"] = v

                samples[well_id] = sample

    return {"sampling_date": global_date, "samples": samples}


# ──────────────────────────────────────────────────────────────────────────────
# MANN-KENDALL
# ──────────────────────────────────────────────────────────────────────────────

def _xl_date(xl_num):
    if not xl_num:
        return None
    try:
        return datetime.fromordinal(datetime(1899, 12, 30).toordinal() + int(xl_num))
    except Exception:
        return None


def _read_mann_kendall(xls_path: str) -> dict:
    wb = xlrd.open_workbook(xls_path)
    result = {}
    for sh_name in wb.sheet_names():
        if sh_name == "Kendall_Dist":
            continue
        ws = wb.sheet_by_name(sh_name)
        well_row = ws.row_values(13)
        wells = [str(w).strip() for w in well_row[3:] if str(w).strip()]
        events = []
        for r in range(14, ws.nrows):
            row = ws.row_values(r)
            if not row[1] or not row[2]:
                break
            dt = _xl_date(row[2])
            if dt is None:
                break
            values = {
                well: (float(row[3 + wi]) if row[3 + wi] != "" else None)
                for wi, well in enumerate(wells)
            }
            events.append((dt, values))
        result[sh_name] = {"wells": wells, "events": events}
    return result


def _update_mann_kendall(mk_data: dict, new_dt: datetime, new_values: dict) -> dict:
    for heb, sh in {"בנזן": "בנזן", "MTBE": "MTBE"}.items():
        if sh not in mk_data:
            continue
        sheet = mk_data[sh]
        well_vals = new_values.get(heb, {})
        sheet["events"].append((new_dt, {w: well_vals.get(w) for w in sheet["wells"]}))
    return mk_data


def _mk_s(values):
    clean = [v for v in values if v is not None]
    n = len(clean)
    if n < 4:
        return 0, n
    s = sum(
        1 if clean[j] > clean[i] else (-1 if clean[j] < clean[i] else 0)
        for i in range(n - 1)
        for j in range(i + 1, n)
    )
    return s, n


def _mk_conf(s, n):
    if n < 4:
        return 0.5
    var_s = n * (n - 1) * (2 * n + 5) / 18
    if var_s == 0:
        return 0.5
    z = (s - 1) / np.sqrt(var_s) if s > 0 else (s + 1) / np.sqrt(var_s) if s < 0 else 0
    try:
        from scipy.special import erfc
        return float(0.5 * erfc(-z / np.sqrt(2)))
    except ImportError:
        return 0.5


def _mk_trend(s, conf, cov):
    if conf >= 0.95:
        return "Increasing" if s > 0 else "Decreasing"
    if conf >= 0.90:
        return "Probably Increasing" if s > 0 else "Probably Decreasing"
    if s > 0:
        return "No Trend"
    return "Stable" if cov < 1.0 else "No Trend"


def _write_mann_kendall(mk_data: dict, orig_xls: str, out_path: str):
    orig_wb = xlrd.open_workbook(orig_xls, formatting_info=True)
    new_wb = xlwt.Workbook()
    for sh_name in orig_wb.sheet_names():
        orig_ws = orig_wb.sheet_by_name(sh_name)
        new_ws = new_wb.add_sheet(sh_name, cell_overwrite_ok=True)
        for r in range(orig_ws.nrows):
            for c in range(orig_ws.ncols):
                new_ws.write(r, c, orig_ws.cell_value(r, c))
        if sh_name == "Kendall_Dist" or sh_name not in mk_data:
            continue
        sheet = mk_data[sh_name]
        wells, events = sheet["wells"], sheet["events"]
        for ei, (dt, vals) in enumerate(events):
            row = 14 + ei
            new_ws.write(row, 1, ei + 1)
            new_ws.write(row, 2, (dt - datetime(1899, 12, 30)).days)
            for wi, w in enumerate(wells):
                v = vals.get(w)
                new_ws.write(row, 3 + wi, v if v is not None else "")
        # Update statistics rows 54-57
        for wi, w in enumerate(wells):
            col = 3 + wi
            all_v = [e[1].get(w) for e in events]
            numeric = [v for v in all_v if v is not None]
            cov = (np.std(numeric) / np.mean(numeric)
                   if numeric and np.mean(numeric) != 0 else 0)
            s, n = _mk_s(all_v)
            conf = _mk_conf(s, n)
            new_ws.write(54, col, round(cov, 4))
            new_ws.write(55, col, s)
            new_ws.write(56, col, round(conf, 4))
            new_ws.write(57, col, _mk_trend(s, conf, cov))
        new_ws.write(3, 2, (datetime.today() - datetime(1899, 12, 30)).days)
    new_wb.save(out_path)


# ──────────────────────────────────────────────────────────────────────────────
# CHART GENERATION
# ──────────────────────────────────────────────────────────────────────────────

PARAM_INFO = {
    "בנזן":      ("Benzene / בנזן",           "mg/L", 0.005),
    "MTBE":      ("MTBE",                      "mg/L", 0.04),
    "טולואן":    ("Toluene / טולואן",          "mg/L", 0.70),
    "אתיל בנזן": ("Ethyl Benzene / אתיל בנזן", "mg/L", 0.30),
    "קסילן":     ("Xylene / קסילן",            "mg/L", 0.50),
}
_COLORS = [
    "#1f77b4", "#ff7f0e", "#2ca02c", "#d62728",
    "#9467bd", "#8c564b", "#e377c2", "#bcbd22",
]


def _generate_charts(mk_data: dict, out_dir: Path) -> dict:
    out_dir.mkdir(parents=True, exist_ok=True)
    paths = {}
    plt.rcParams.update({"font.family": "DejaVu Sans", "axes.unicode_minus": False})

    for sh_name, sheet in mk_data.items():
        wells, events = sheet["wells"], sheet["events"]
        if not events:
            continue
        label, unit, thresh = PARAM_INFO.get(sh_name, (sh_name, "mg/L", None))
        dates = [e[0] for e in events]

        fig, ax = plt.subplots(figsize=(10, 5))
        plotted = False
        for wi, well in enumerate(wells):
            vals = [e[1].get(well) for e in events]
            pairs = [(d, v) for d, v in zip(dates, vals) if v is not None]
            if not pairs:
                continue
            xs, ys = zip(*pairs)
            ax.plot(xs, ys, marker="o", label=well,
                    color=_COLORS[wi % len(_COLORS)], linewidth=1.5, markersize=5)
            plotted = True

        if not plotted:
            plt.close(fig)
            continue

        if thresh:
            ax.axhline(thresh, color="red", linestyle="--", linewidth=1.2,
                       label=f"Target ({thresh} {unit})")

        ax.xaxis.set_major_formatter(mdates.DateFormatter("%m/%y"))
        ax.xaxis.set_major_locator(mdates.AutoDateLocator())
        fig.autofmt_xdate()
        ax.set_xlabel("תאריך", fontsize=10)
        ax.set_ylabel(unit, fontsize=10)
        ax.set_title(label, fontsize=12, fontweight="bold")
        ax.legend(loc="upper right", fontsize=8, ncol=2)
        ax.grid(True, alpha=0.3)
        ax.text(0.01, 0.01,
                "ריכוזים קטנים מ-0.01 מג/ל מוצגים בתחתית הגרף",
                transform=ax.transAxes, fontsize=7, color="gray", va="bottom")
        plt.tight_layout()
        p = out_dir / f"chart_{sh_name}.png"
        plt.savefig(p, dpi=150, bbox_inches="tight")
        plt.close(fig)
        paths[sh_name] = str(p)

    return paths


# ──────────────────────────────────────────────────────────────────────────────
# WORD DOCUMENT UPDATER
# ──────────────────────────────────────────────────────────────────────────────

THRESH = {
    "בנזן": 0.005, "טולואן": 0.70, "אתיל בנזן": 0.30,
    "קסילן": 0.50, "MTBE": 0.04,
}
COL_PARAMS = ["MTBE", "בנזן", "טולואן", "אתיל בנזן", "קסילן"]


def _set_cell(cell, text, bold=False, color=None):
    """Replace cell content with a single run."""
    for para in cell.paragraphs:
        for child in list(para._element):
            if child.tag.endswith("}r"):
                para._element.remove(child)
    run = cell.paragraphs[0].add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = color


def _insert_row_after(table, after_idx: int, tmpl_idx: int):
    """Insert a deep-copied row after after_idx."""
    tmpl_tr = copy.deepcopy(table.rows[tmpl_idx]._tr)
    for t in tmpl_tr.findall(".//" + qn("w:t")):
        t.text = ""
    table.rows[after_idx]._tr.addnext(tmpl_tr)
    return table.rows[after_idx + 1]


def _last_row_for_well(table, well_id: str) -> int:
    last = -1
    for i, row in enumerate(table.rows):
        if row.cells[0].text.strip() == well_id:
            last = i
    return last


def _update_field_table(doc: Document, field_data: dict, wells: list):
    tbl = doc.tables[1]
    header = tbl.rows[0]
    well_col = {
        c.text.strip(): ci
        for ci, c in enumerate(header.cells)
        if c.text.strip() in wells
    }
    FIELD_ROWS = {
        "pH": 1, "EC": 2, "טמפרטורה": 3, "חמצן מומס": 4,
        "עכירות": 5, "רדוקס": 6, "עומק פני המים": 7,
        "עומק כללי של הקידוח": 8, "עומק דגימה מפני המים": 9,
    }
    for field, ri in FIELD_ROWS.items():
        if ri >= len(tbl.rows):
            continue
        row = tbl.rows[ri]
        for well, ci in well_col.items():
            val = field_data.get(well, {}).get("field", {}).get(field)
            if val is not None:
                txt = str(int(val)) if isinstance(val, float) and val == int(val) else str(val)
                _set_cell(row.cells[ci], txt)


def _update_historical_tables(doc: Document, new_date: str,
                               samples: dict, wells_order: list):
    # Map well → table index
    well_table = {}
    for ti, tbl in enumerate(doc.tables[2:], 2):
        for row in tbl.rows:
            w = row.cells[0].text.strip()
            if w in wells_order:
                well_table[w] = ti

    for well in wells_order:
        if well not in samples:
            continue
        ti = well_table.get(well)
        if ti is None:
            continue
        tbl = doc.tables[ti]
        sample = samples[well]
        results = sample.get("results", {})
        floating = sample.get("floating_layer", False)

        last = _last_row_for_well(tbl, well)
        if last < 0:
            continue

        new_row = _insert_row_after(tbl, last, last)
        cells = new_row.cells
        # NOTE: cells[0] (well name) is often a vertically-merged cell shared
        # across the whole table — do NOT write to it here, or it will wipe
        # out the merge-origin text for every row in the table.
        _set_cell(cells[1], new_date)
        _set_cell(cells[2], "")

        if floating:
            _set_cell(cells[3], "לא נדגם עקב שכבה צפה")
            for ci in range(4, min(8, len(cells))):
                _set_cell(cells[ci], "")
        else:
            for ci, param in enumerate(COL_PARAMS):
                if 3 + ci >= len(cells):
                    break
                val = results.get(param)
                thresh = THRESH.get(param)
                if val is None:
                    txt, exceeds = "<0.001", False
                else:
                    exceeds = thresh is not None and val > thresh
                    txt = str(val)
                color = RGBColor(0xC0, 0x00, 0x00) if exceeds else None
                _set_cell(cells[3 + ci], txt, bold=exceeds, color=color)


def _replace_chart_image(doc: Document, shape_idx: int, img_path: str):
    if shape_idx >= len(doc.inline_shapes):
        return
    shape = doc.inline_shapes[shape_idx]
    w_emu, h_emu = shape.width, shape.height
    blip = shape._inline.find(".//" + qn("a:blip"))
    if blip is None:
        return
    rId = blip.get(qn("r:embed"))
    image_part = doc.part.related_parts[rId]
    with open(img_path, "rb") as f:
        image_part._blob = f.read()
    extent = shape._inline.find(qn("wp:extent"))
    if extent is not None:
        extent.set("cx", str(w_emu))
        extent.set("cy", str(h_emu))


# ──────────────────────────────────────────────────────────────────────────────
# PUBLIC API
# ──────────────────────────────────────────────────────────────────────────────

def run_update_bytes(
    word_bytes: bytes,
    lab_pdf_bytes,
    mk_xls_bytes: bytes,
    field_pdf_bytes: bytes = None,   # reserved for future OCR use
    lab_type: str = "bactochem",     # "bactochem" or "aminolab"
) -> tuple:
    """
    Process all inputs in memory and return (updated_word_bytes, updated_mk_xls_bytes).

    lab_pdf_bytes:
      - "bactochem": a single PDF's bytes (one PDF covers all wells).
      - "aminolab":  bytes of ONE PDF, or a list/tuple of bytes — one per
        well/certificate — since each Aminolab certificate covers a single well.
    """
    import os

    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)

        # Write inputs to temp files
        word_path = tmp / "report.docx"
        mk_path   = tmp / "mk.xls"
        out_word  = tmp / "updated_report.docx"
        out_mk    = tmp / "updated_mk.xls"

        word_path.write_bytes(word_bytes)
        mk_path.write_bytes(mk_xls_bytes)

        # ── Parse lab results ──────────────────────────────────────
        if lab_type == "aminolab":
            pdf_bytes_list = (
                lab_pdf_bytes if isinstance(lab_pdf_bytes, (list, tuple))
                else [lab_pdf_bytes]
            )
            lab_paths = []
            for i, pdf_bytes in enumerate(pdf_bytes_list):
                p = tmp / f"lab_{i}.pdf"
                p.write_bytes(pdf_bytes)
                lab_paths.append(str(p))
            lab = parse_aminolab_pdfs(lab_paths)
            samples, new_date = lab["samples"], lab["sampling_date"]
            if not new_date:
                raise ValueError("לא ניתן לחלץ תאריך דיגום מה-PDF. ודא שהקובץ הוא תעודת אמינולאב תקנית.")
            if not samples:
                raise ValueError("לא נמצאו דגימות ב-PDF. ודא שהקובץ הוא תעודת אמינולאב לדיגום מי תהום.")
        else:
            lab_path = tmp / "lab.pdf"
            lab_path.write_bytes(lab_pdf_bytes)
            lab = parse_bactochem_pdf(str(lab_path))
            samples, new_date = lab["samples"], lab["sampling_date"]
            if not new_date:
                raise ValueError("לא ניתן לחלץ תאריך דיגום מה-PDF. ודא שהקובץ הוא דוח בקטוכם תקני.")
            if not samples:
                raise ValueError("לא נמצאו דגימות ב-PDF. ודא שהקובץ הוא דוח בקטוכם לדיגום מי תהום.")

        # ── Update Mann-Kendall ────────────────────────────────────
        mk_data = _read_mann_kendall(str(mk_path))
        new_dt = datetime.strptime(new_date, "%d.%m.%y")
        mk_new = {}
        for well, s in samples.items():
            for param, val in s.get("results", {}).items():
                mk_new.setdefault(param, {})[well] = val
        mk_data = _update_mann_kendall(mk_data, new_dt, mk_new)

        # ── Generate charts ────────────────────────────────────────
        chart_paths = _generate_charts(mk_data, tmp / "charts")

        # ── Write updated MK XLS ──────────────────────────────────
        _write_mann_kendall(mk_data, str(mk_path), str(out_mk))

        # ── Update Word document ──────────────────────────────────
        doc = Document(str(word_path))

        # Detect the well-naming convention actually used in THIS document
        # (different sites use different prefixes, e.g. "מת-1" or "צא - 1"),
        # then remap the parsed sample keys (always "מת-N") onto it by well number.
        doc_wells = set()
        for tbl in doc.tables:
            for row in tbl.rows:
                if not row.cells:
                    continue
                w = row.cells[0].text.strip()
                if w and re.search(r'\d', w) and len(w) < 20:
                    doc_wells.add(w)

        num_to_doc_well = {}
        for w in doc_wells:
            m = re.search(r'(\d+)\s*$', w)
            if m:
                num_to_doc_well.setdefault(m.group(1), w)

        remapped_samples = {}
        unmatched = []
        for well_key, sample in samples.items():
            m = re.search(r'(\d+)$', well_key)
            num = m.group(1) if m else None
            doc_well = num_to_doc_well.get(num)
            if doc_well:
                remapped_samples[doc_well] = sample
            else:
                remapped_samples[well_key] = sample
                unmatched.append(well_key)
        samples = remapped_samples
        wells_order = list(samples.keys())

        if unmatched:
            raise ValueError(
                "לא נמצאה בדוח ה-Word התאמה לבארות הבאות מה-PDF: "
                + ", ".join(unmatched)
                + ". ודא שמספרי הבארות בדוח ה-PDF תואמים לבארות בדוח ה-Word."
            )

        _update_field_table(doc, samples, wells_order)
        _update_historical_tables(doc, new_date, samples, wells_order)

        # Replace chart images (shape 1 = Benzene, shape 2 = MTBE)
        for chart_idx, sh_name in enumerate(["בנזן", "MTBE"]):
            if sh_name in chart_paths:
                _replace_chart_image(doc, 1 + chart_idx, chart_paths[sh_name])

        doc.save(str(out_word))

        return out_word.read_bytes(), out_mk.read_bytes()
