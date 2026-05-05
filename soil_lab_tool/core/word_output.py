"""
core/word_output.py
--------------------
Builds a multi-section Hebrew RTL Word lab report using python-docx.

One section per analysis type found in the records.

Layout (compounds as rows, samples as columns):
  Col 0 : תרכובת
  Col 1 : CAS Number
  Col 2…: threshold column(s)
  next  : יחידות
  rest  : one column per sample (header = borehole\\ndepth)

Colour coding (cell background):
  Yellow  (FFFF00) — detected value exceeds a VSL threshold
                     (VSL_SOIL, GW, PFAS_VSL)
  Orange  (FFC000) — detected value exceeds a Tier-1 threshold
                     (any TIER1_*, GAS_*, PFAS_TIER1_* key)
  Gray    (D3D3D3) — non-detect / below-limit but LOD > strictest threshold

Standalone helpers (add a single section to an existing Document):
  add_tph_table(doc, records, tm, selected_thresholds)
  add_metals_table(doc, records, tm, selected_thresholds)
  add_voc_table(doc, records, tm, selected_thresholds)
  add_pfas_table(doc, records, tm, selected_thresholds)

Main class:
  LabReportWord – mirrors the LabReportExcel API
"""

from __future__ import annotations

import io
import os
import re
from collections import defaultdict
from datetime import date

import pandas as pd
from openpyxl import load_workbook

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Twips

from core.threshold_manager import (
    ThresholdManager,
    ANALYSIS_THRESHOLDS,
    THRESHOLD_LABELS,
)
from core.excel_output import SHEET_CONFIG, _ordered_unique, _split_sample_depth

# ── Colour constants ──────────────────────────────────────────────────
_YELLOW = "FFFF00"   # VSL exceedance
_ORANGE = "FFC000"   # Tier-1 exceedance
_GRAY   = "D3D3D3"   # LOD > threshold (uncertain non-detect)

# Threshold keys that yield yellow on exceedance; all others yield orange
_VSL_KEYS: frozenset[str] = frozenset({"VSL_SOIL", "GW", "PFAS_VSL"})

# Font names / sizes
_FONT_HE = "David"
_FONT_EN = "Times New Roman"
_SZ_HE   = Pt(9)
_SZ_EN   = Pt(8)
_SZ_SM   = Pt(8)   # data cells


# ── Low-level XML / formatting helpers ────────────────────────────────

def _has_hebrew(s: str) -> bool:
    return any('א' <= c <= 'ת' for c in str(s))


def _pick_font(s: str) -> tuple[str, Pt]:
    """Return (font_name, font_size) based on whether content is Hebrew."""
    s = str(s) if s is not None else ""
    has_he = _has_hebrew(s)
    has_en = any(c.isalpha() and c.isascii() for c in s)
    if has_he or not has_en:
        return _FONT_HE, _SZ_HE
    return _FONT_EN, _SZ_EN


def _set_cell_bg(cell, hex_color: str) -> None:
    """Set table cell background to a solid hex colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for ex in tcPr.findall(qn("w:shd")):
        tcPr.remove(ex)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_rtl_para(para) -> None:
    """Enable RTL text direction on a paragraph."""
    pPr  = para._p.get_or_add_pPr()
    bidi = pPr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        pPr.append(bidi)
    bidi.set(qn("w:val"), "1")


def _set_rtl_table(table) -> None:
    """Set bidiVisual on a table so columns display right-to-left."""
    tblPr = table._tbl.get_or_add_tblPr()
    if tblPr.find(qn("w:bidiVisual")) is None:
        tblPr.append(OxmlElement("w:bidiVisual"))


def _set_table_full_width(table) -> None:
    """Stretch table to 100 % of the text area."""
    tblPr = table._tbl.get_or_add_tblPr()
    for ex in tblPr.findall(qn("w:tblW")):
        tblPr.remove(ex)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "pct")
    tblW.set(qn("w:w"),    "5000")   # 5000 = 100 % in fiftieths of a percent
    tblPr.append(tblW)


def _clear_para_runs(para) -> None:
    p = para._p
    for r in p.findall(qn("w:r")):
        p.remove(r)


def _write_cell(
    cell,
    text: str,
    bold: bool = False,
    italic: bool = False,
    size: Pt | None = None,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER,
    gray_text: bool = False,
) -> None:
    """Write *text* into *cell* with auto font selection and RTL support."""
    para = cell.paragraphs[0]
    _clear_para_runs(para)
    _set_rtl_para(para)
    para.alignment = align
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if text is None or str(text) == "":
        return

    run = para.add_run(str(text))
    fname, fsize = _pick_font(str(text))
    run.font.name   = fname
    run.font.size   = size or fsize
    run.font.bold   = bold
    run.font.italic = italic
    if gray_text:
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


# ── Number / value formatting ─────────────────────────────────────────

def _fmt_lod(lod: float) -> str:
    if lod == int(lod):
        return str(int(lod))
    return f"{round(lod, 3):.3f}".rstrip("0").rstrip(".")


def _fmt_thresh(val) -> str:
    if val is None:
        return "לא קיים"
    if isinstance(val, (int, float)):
        if val == int(val):
            return str(int(val))
        return f"{round(val, 2):.2f}".rstrip("0").rstrip(".")
    return str(val)


def _fmt_value(value, flag: str, lod) -> str:
    """Format a measured sample value for Word display."""
    if flag == "ND" or (value is None and flag not in ("<LOD", "<LOQ", "<")):
        return _fmt_lod(lod) if isinstance(lod, (int, float)) else "ND"
    if flag == "<LOD":
        return f"<{_fmt_lod(lod)}" if isinstance(lod, (int, float)) else "ND"
    if flag == "<LOQ":
        ref = lod if isinstance(lod, (int, float)) else value
        return _fmt_lod(ref) if isinstance(ref, (int, float)) else "ND"
    if flag == "<":
        return f"<{round(value, 2)}" if isinstance(value, (int, float)) else f"<{value}"
    if isinstance(value, float):
        return f"{round(value, 3):.3f}".rstrip("0").rstrip(".")
    return str(value) if value is not None else "ND"


# ── Colour-coding logic ───────────────────────────────────────────────

def _cell_color(
    value,
    flag: str,
    lod,
    t_vals: dict,
    thresh_keys: list[str],
) -> str | None:
    """
    Return hex background colour for a sample cell, or None.

    Priority: Orange (Tier-1) > Yellow (VSL) > Gray (LOD > strictest threshold).
    """
    is_nondetect = flag in ("ND", "<LOD", "<LOQ", "<")

    if is_nondetect:
        lod_ref: float | None = None
        if isinstance(lod, (int, float)):
            lod_ref = lod
        elif isinstance(value, (int, float)) and flag in ("<", "<LOQ"):
            lod_ref = value
        if lod_ref is not None:
            strictest = min(
                (v for v in t_vals.values() if isinstance(v, (int, float))),
                default=None,
            )
            if strictest is not None and lod_ref > strictest:
                return _GRAY
        return None

    if not isinstance(value, (int, float)):
        return None

    # Tier-1 → orange (checked first for priority)
    for k in thresh_keys:
        if k not in _VSL_KEYS:
            t = t_vals.get(k)
            if isinstance(t, (int, float)) and value > t:
                return _ORANGE

    # VSL → yellow
    for k in thresh_keys:
        if k in _VSL_KEYS:
            t = t_vals.get(k)
            if isinstance(t, (int, float)) and value > t:
                return _YELLOW

    return None


# ── Document-level helpers ────────────────────────────────────────────

def _add_section_heading(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    _set_rtl_para(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(title)
    run.font.name = _FONT_HE
    run.font.size = Pt(12)
    run.font.bold = True


def _add_legend(doc: Document, include_gray: bool = False) -> None:
    """Append a colour legend below a table."""
    items: list[tuple[str, str]] = [
        ("חריגה מערך סף VSL",    _YELLOW),
        ("חריגה מערך סף Tier-1", _ORANGE),
    ]
    if include_gray:
        items.append(("ספג גילוי גבוה מערך הסף", _GRAY))

    for label, hex_color in items:
        p = doc.add_paragraph()
        _set_rtl_para(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f" {label} ")
        run.font.name = _FONT_HE
        run.font.size = _SZ_HE
        run.font.bold = True
        # Background colour applied as character shading (w:rPr > w:shd)
        rPr = run._r.get_or_add_rPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        rPr.append(shd)


def _ordered_thresh_keys(atype: str, sel: list[str] | None) -> list[str]:
    valid = set(ANALYSIS_THRESHOLDS.get(atype, []))
    if sel is not None:
        return [k for k in sel if k in valid]
    return ANALYSIS_THRESHOLDS.get(atype, [])


# ── Core table builder ────────────────────────────────────────────────

def _build_analysis_table(
    doc: Document,
    records: list[dict],
    tm: ThresholdManager,
    cfg: dict,
    thresh_keys: list[str],
) -> None:
    """
    Append one analysis-type table (compounds × samples) to *doc*.

    Header columns displayed right-to-left:
      תרכובת | CAS Number | threshold(s) | יחידות | sample_1 … sample_n
    """
    samples   = _ordered_unique(r["sample_id"] for r in records)
    compounds = _ordered_unique(r["compound"]  for r in records)
    unit      = cfg.get("unit", "")

    if not compounds or not samples:
        return

    # Pivot: compound → sample_id → (value, flag, lod)
    pivot:    dict[str, dict] = {}
    cas_map:  dict[str, str]  = {}
    unit_map: dict[str, str]  = {}
    for r in records:
        cmp = r["compound"]
        sid = r["sample_id"]
        if cmp not in pivot:
            pivot[cmp]    = {}
            cas_map[cmp]  = r.get("cas", "")
            unit_map[cmp] = r.get("unit", unit)
        pivot[cmp][sid] = (r.get("value"), r.get("flag", ""), r.get("lod"))

    # Threshold values per compound
    thresh_vals: dict[str, dict] = {
        cmp: {
            k: tm.get_threshold_with_name(cas_map[cmp], k, compound_name=cmp)
            for k in thresh_keys
        }
        for cmp in compounds
    }

    split_map = {sid: _split_sample_depth(sid) for sid in samples}

    n_thresh = len(thresh_keys)
    # Fixed cols: תרכובת(0) | CAS(1) | thresh…(2…n+1) | יחידות(n+2)
    n_fixed  = 2 + n_thresh + 1
    n_cols   = n_fixed + len(samples)

    # ── Create table ──────────────────────────────────────────────────
    table = doc.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"
    _set_rtl_table(table)
    _set_table_full_width(table)

    # ── Header row ────────────────────────────────────────────────────
    hdr = table.rows[0].cells
    _write_cell(hdr[0], "תרכובת",     bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _write_cell(hdr[1], "CAS Number",  bold=True)
    for i, k in enumerate(thresh_keys):
        _write_cell(hdr[2 + i], THRESHOLD_LABELS.get(k, k), bold=True)
    _write_cell(hdr[2 + n_thresh], "יחידות", bold=True)
    for j, sid in enumerate(samples):
        bh, dep = split_map[sid]
        hdr_text = f"{bh}\n{dep}" if dep else bh
        _write_cell(hdr[n_fixed + j], hdr_text, bold=True)

    # ── Data rows ─────────────────────────────────────────────────────
    has_gray = False

    for cmp in compounds:
        row_cells = table.add_row().cells
        cas    = cas_map.get(cmp, "")
        t_vals = thresh_vals.get(cmp, {})
        u      = unit_map.get(cmp, unit)

        _write_cell(row_cells[0], cmp, align=WD_ALIGN_PARAGRAPH.RIGHT)
        _write_cell(row_cells[1], cas if cas else "—", size=_SZ_SM)

        for i, k in enumerate(thresh_keys):
            tv   = t_vals.get(k)
            tv_r = round(tv, 2) if isinstance(tv, (int, float)) else None
            text = _fmt_thresh(tv_r)
            _write_cell(row_cells[2 + i], text,
                        italic=   (text == "לא קיים"),
                        gray_text=(text == "לא קיים"),
                        size=_SZ_SM)

        _write_cell(row_cells[2 + n_thresh], u, size=_SZ_SM)

        for j, sid in enumerate(samples):
            v, flag, lod = pivot.get(cmp, {}).get(sid, (None, "ND", None))
            display = _fmt_value(v, flag, lod)
            cell    = row_cells[n_fixed + j]

            color = _cell_color(v, flag, lod, t_vals, thresh_keys)
            if color:
                _set_cell_bg(cell, color)
                if color == _GRAY:
                    has_gray = True

            bold_val = color in (_ORANGE, _YELLOW)
            _write_cell(cell, display, bold=bold_val, size=_SZ_SM)

    doc.add_paragraph()
    _add_legend(doc, include_gray=has_gray)
    doc.add_paragraph()


# ── Standalone section-adder functions ───────────────────────────────

def add_tph_table(
    doc: Document,
    records: list[dict],
    tm: ThresholdManager,
    selected_thresholds: list[str] | None = None,
) -> None:
    """Append a TPH analysis section (heading + table + legend) to *doc*."""
    atype = "SOIL_TPH"
    recs  = [r for r in records if r.get("analysis_type") == atype]
    if not recs:
        return
    cfg         = SHEET_CONFIG.get(atype, {"name": "קרקע TPH", "unit": "mg/kg"})
    thresh_keys = _ordered_thresh_keys(atype, selected_thresholds)
    _add_section_heading(doc, cfg["name"])
    _build_analysis_table(doc, recs, tm, cfg, thresh_keys)


def add_metals_table(
    doc: Document,
    records: list[dict],
    tm: ThresholdManager,
    selected_thresholds: list[str] | None = None,
) -> None:
    """Append a Metals analysis section to *doc*."""
    atype = "SOIL_METALS"
    recs  = [r for r in records if r.get("analysis_type") == atype]
    if not recs:
        return
    cfg         = SHEET_CONFIG.get(atype, {"name": "קרקע מתכות", "unit": "mg/kg DW"})
    thresh_keys = _ordered_thresh_keys(atype, selected_thresholds)
    _add_section_heading(doc, cfg["name"])
    _build_analysis_table(doc, recs, tm, cfg, thresh_keys)


def add_voc_table(
    doc: Document,
    records: list[dict],
    tm: ThresholdManager,
    selected_thresholds: list[str] | None = None,
) -> None:
    """
    Append VOC/BTEX sections to *doc*.

    Covers SOIL_VOC, SOIL_MBTEX, GW_VOC, SOIL_GAS_VOC — one subsection each.
    """
    for atype in ("SOIL_VOC", "SOIL_MBTEX", "GW_VOC", "SOIL_GAS_VOC"):
        recs = [r for r in records if r.get("analysis_type") == atype]
        if not recs:
            continue
        cfg         = SHEET_CONFIG.get(atype, {"name": atype, "unit": ""})
        thresh_keys = _ordered_thresh_keys(atype, selected_thresholds)
        _add_section_heading(doc, cfg["name"])
        _build_analysis_table(doc, recs, tm, cfg, thresh_keys)


def add_pfas_table(
    doc: Document,
    records: list[dict],
    tm: ThresholdManager,
    selected_thresholds: list[str] | None = None,
) -> None:
    """
    Append PFAS sections to *doc*.

    Covers SOIL_PFAS and GW_PFAS — one subsection each.
    """
    for atype in ("SOIL_PFAS", "GW_PFAS"):
        recs = [r for r in records if r.get("analysis_type") == atype]
        if not recs:
            continue
        cfg         = SHEET_CONFIG.get(atype, {"name": atype, "unit": ""})
        thresh_keys = _ordered_thresh_keys(atype, selected_thresholds)
        _add_section_heading(doc, cfg["name"])
        _build_analysis_table(doc, recs, tm, cfg, thresh_keys)


# ── Main class ────────────────────────────────────────────────────────

class LabReportWord:
    """
    Build a multi-section Hebrew RTL Word lab report.

    Parameters mirror LabReportExcel for a consistent API.

    Parameters
    ----------
    records : list[dict]
        Flat list of measurement records, each with keys:
        compound, cas, sample_id, value, flag, unit, lod, analysis_type
    threshold_manager : ThresholdManager
    output_path : str | BytesIO
    project_name : str
    client : str
    report_date : str   (DD.MM.YYYY)
    selected_thresholds : list[str] | None
        Override which threshold keys to show.  None = defaults per analysis type.
    combine_tph_voc : bool
        Merge SOIL_TPH + SOIL_VOC records into one combined table.
    combine_tph_mbtex : bool
        Merge SOIL_TPH + SOIL_MBTEX records into one combined table.
    """

    def __init__(
        self,
        records: list[dict],
        threshold_manager: ThresholdManager,
        output_path: str = "lab_report.docx",
        project_name: str = "",
        client: str = "",
        report_date: str = "",
        selected_thresholds: list[str] | None = None,
        combine_tph_voc: bool = False,
        combine_tph_mbtex: bool = False,
    ):
        self.records           = records
        self.tm                = threshold_manager
        self.out_path          = output_path
        self.project           = project_name
        self.client            = client
        self.rep_date          = report_date or date.today().strftime("%d.%m.%Y")
        self.sel_thresh        = selected_thresholds
        self.combine_tph_voc   = combine_tph_voc
        self.combine_tph_mbtex = combine_tph_mbtex

    # ------------------------------------------------------------------
    def build(self) -> str:
        """Generate the Word document and return output_path."""
        doc = Document()

        # Landscape A4 with narrow margins to accommodate wide tables
        section = doc.sections[0]
        section.orientation   = WD_ORIENT.LANDSCAPE
        section.page_width    = Cm(29.7)
        section.page_height   = Cm(21.0)
        section.left_margin   = Cm(1.5)
        section.right_margin  = Cm(1.5)
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)

        self._write_header(doc)

        # Group records by analysis_type
        groups: dict[str, list[dict]] = defaultdict(list)
        for r in self.records:
            groups[r.get("analysis_type", "UNKNOWN")].append(r)

        if self.combine_tph_voc and "SOIL_TPH" in groups and "SOIL_VOC" in groups:
            groups["SOIL_TPH_VOC"] = (
                list(groups.pop("SOIL_TPH")) + list(groups.pop("SOIL_VOC"))
            )

        if self.combine_tph_mbtex and "SOIL_TPH" in groups and "SOIL_MBTEX" in groups:
            groups["SOIL_TPH_MBTEX"] = (
                list(groups.pop("SOIL_TPH")) + list(groups.pop("SOIL_MBTEX"))
            )

        _ORDER = [
            "SOIL_TPH", "SOIL_TPH_VOC", "SOIL_TPH_MBTEX",
            "SOIL_METALS",
            "SOIL_VOC", "SOIL_MBTEX",
            "SOIL_PFAS",
            "GW_VOC", "GW_PFAS",
            "SOIL_GAS_VOC",
        ]
        ordered = [k for k in _ORDER if k in groups]
        ordered += [k for k in groups if k not in ordered and k != "LOWFLOW"]

        for atype in ordered:
            recs        = groups[atype]
            cfg         = SHEET_CONFIG.get(atype, {"name": atype, "unit": ""})
            thresh_keys = _ordered_thresh_keys(atype, self.sel_thresh)
            _add_section_heading(doc, cfg["name"])
            _build_analysis_table(doc, recs, self.tm, cfg, thresh_keys)

        if isinstance(self.out_path, (str, os.PathLike)):
            os.makedirs(os.path.dirname(str(self.out_path)) or ".", exist_ok=True)
        doc.save(self.out_path)
        return self.out_path

    # ------------------------------------------------------------------
    def _write_header(self, doc: Document) -> None:
        p = doc.add_paragraph()
        _set_rtl_para(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        text = "   |   ".join([
            f"שם פרויקט: {self.project}",
            f"תאריך: {self.rep_date}",
            f"מזמין: {self.client}",
        ])
        run = p.add_run(text)
        run.font.name = _FONT_HE
        run.font.size = Pt(11)
        run.font.bold = True
        doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# ALS-based Word builder — Ofri's standalone pipeline
# Public API: parse_als_file, build_word_report, load_threshold_file,
#             get_tier1_col, tier1_label
# ══════════════════════════════════════════════════════════════════════════════

# ── Page / colour constants ───────────────────────────────────────
_ALS_COLOR_YELLOW   = "FFFF00"
_ALS_COLOR_ORANGE   = "FFC000"
_ALS_COLOR_HEADER   = "B7D7F0"
_ALS_COLOR_HEADER2  = "00B0F0"
_ALS_COLOR_WHITE    = "FFFFFF"

_ALS_PAGE_SIZES = {
    "A4":      (11906, 16838),
    "Tabloid": (17280, 22320),
}
_ALS_MARGIN = 540   # narrow margins in twips

# ── Compound maps & aliases ───────────────────────────────────────
_ALS_METAL_MAP = {
    "aluminium":"Al","aluminum":"Al","antimony":"Sb","arsenic":"As",
    "barium":"Ba","beryllium":"Be","bismuth":"Bi","boron":"B",
    "cadmium":"Cd","calcium":"Ca","chromium":"Cr","cobalt":"Co",
    "copper":"Cu","iron":"Fe","lead":"Pb","lithium":"Li",
    "magnesium":"Mg","manganese":"Mn","mercury":"Hg","nickel":"Ni",
    "potassium":"K","selenium":"Se","silver":"Ag","sodium":"Na",
    "vanadium":"V","zinc":"Zn","molybdenum":"Mo","tin":"Sn",
    "titanium":"Ti","strontium":"Sr","thallium":"Tl",
    "phosphorus":"P","sulphur":"S","silicon":"Si",
}

_ALS_METALS_ORDER = [
    "Al","Sb","As","Ba","Be","Bi","B","Cd","Ca","Cr","Co","Cu","Fe",
    "Pb","Li","Mg","Mn","Hg","Ni","K","Se","Ag","Na","V","Zn",
]

_ALS_THRESH_METAL_MAP = {
    "aluminum":"Al","antimony (metallic)":"Sb","antimony":"Sb",
    "arsenic, inorganic":"As","arsenic":"As","barium":"Ba",
    "beryllium and compounds":"Be","beryllium":"Be",
    "boron and borates only":"B","boron":"B",
    "cadmium (water) source: water and air":"Cd","cadmium":"Cd",
    "calcium":"Ca","chromium, total":"Cr","chromium":"Cr","cobalt":"Co",
    "copper":"Cu","iron":"Fe","lead and compounds":"Pb","lead":"Pb",
    "lithium":"Li","magnesium":"Mg","manganese (non-diet)":"Mn","manganese":"Mn",
    "mercuric chloride (and other mercury salts)":"Hg","mercury":"Hg",
    "nickel soluble salts":"Ni","nickel":"Ni","potassium":"K",
    "selenium":"Se","silver":"Ag","sodium":"Na",
    "vanadium and compounds":"V","vanadium":"V",
    "zinc and compounds":"Zn","zinc":"Zn","molybdenum":"Mo","tin":"Sn",
    "titanium":"Ti","strontium":"Sr","thallium":"Tl",
    "phosphorus":"P","sulphur":"S","silicon":"Si",
}

_ALS_VOC_ALIAS = {
    "1.2.4-trimethylbenzene":           "Trimethylbenzene, 1,2,4-",
    "1,2,4-trimethylbenzene":           "Trimethylbenzene, 1,2,4-",
    "1.3.5-trimethylbenzene":           "Trimethylbenzene, 1,3,5-",
    "1,3,5-trimethylbenzene":           "Trimethylbenzene, 1,3,5-",
    "mtbe":                             "Methyl tert-Butyl Ether (MTBE)",
    "methyl tert-butyl ether":          "Methyl tert-Butyl Ether (MTBE)",
    "n-propylbenzene":                  "Propyl benzene",
    "propylbenzene":                    "Propyl benzene",
    "isopropylbenzene":                 "Cumene",
    "cumene":                           "Cumene",
    "2-butanone (mek)":                 "Methyl Ethyl Ketone - MEK (2-Butanone)",
    "2-butanone":                       "Methyl Ethyl Ketone - MEK (2-Butanone)",
    "methyl ethyl ketone":              "Methyl Ethyl Ketone - MEK (2-Butanone)",
    "mek":                              "Methyl Ethyl Ketone - MEK (2-Butanone)",
    "sum of xylenes":                   "Xylenes",
    "xylenes, total":                   "Xylenes",
    "total xylenes":                    "Xylenes",
    "1.1-dichloroethene":               "1,1-Dichloroethylene",
    "1,1-dichloroethene":               "1,1-Dichloroethylene",
    "1,1-dichloroethylene":             "1,1-Dichloroethylene",
    "1.2-dichloroethane":               "1,2- (EDC) Dichloroethane",
    "1,2-dichloroethane":               "1,2- (EDC) Dichloroethane",
    "ethylene dichloride":              "1,2- (EDC) Dichloroethane",
    "dichloromethane":                  "Methylene Chloride",
    "methylene chloride":               "Methylene Chloride",
    "tetrachloroethene":                "Tetrachloroethylene (PCE)",
    "tetrachloroethylene":              "Tetrachloroethylene (PCE)",
    "perchloroethylene":                "Tetrachloroethylene (PCE)",
    "pce":                              "Tetrachloroethylene (PCE)",
    "tetrachloromethane":               "Carbon Tetrachloride",
    "carbon tetrachloride":             "Carbon Tetrachloride",
    "trichloroethene":                  "Trichloroethylene (TCE)",
    "trichloroethylene":                "Trichloroethylene (TCE)",
    "tce":                              "Trichloroethylene (TCE)",
    "cis-1.2-dichloroethene":           "1,2-cis-Dichloroethylene",
    "cis-1,2-dichloroethene":           "1,2-cis-Dichloroethylene",
    "cis-1,2-dichloroethylene":         "1,2-cis-Dichloroethylene",
    "1,2-cis-dichloroethylene":         "1,2-cis-Dichloroethylene",
    "trans-1.2-dichloroethene":         "1,2-trans-Dichloroethylene",
    "trans-1,2-dichloroethene":         "1,2-trans-Dichloroethylene",
    "trans-1,2-dichloroethylene":       "1,2-trans-Dichloroethylene",
    "1,2-trans-dichloroethylene":       "1,2-trans-Dichloroethylene",
    "benz(a)anthracene":                "Benz[a]anthracene",
    "benz[a]anthracene":                "Benz[a]anthracene",
    "benzo(g.h.i)perylene":             "h,i)perylene Benzo(g",
    "benzo(g,h,i)perylene":             "h,i)perylene Benzo(g",
    "benzo[g,h,i]perylene":             "h,i)perylene Benzo(g",
    "dibenz(a.h)anthracene":            "h]anthracene Dibenz[a",
    "dibenz(a,h)anthracene":            "h]anthracene Dibenz[a",
    "dibenz[a,h]anthracene":            "h]anthracene Dibenz[a",
    "indeno(1.2.3.cd)pyrene":           "2,3-cd]pyrene Indeno[1",
    "indeno(1,2,3-cd)pyrene":           "2,3-cd]pyrene Indeno[1",
    "indeno[1,2,3-cd]pyrene":           "2,3-cd]pyrene Indeno[1",
    "4-chloroaniline":                  "p-Chloroaniline",
    "p-chloroaniline":                  "p-Chloroaniline",
    "1-chloronaphthalene":              "Beta-Chloronaphthalene",
    "2-chloronaphthalene":              "Beta-Chloronaphthalene",
    "6-caprolactam":                    "Caprolactam",
    "2.4.5-trichlorophenol":            "Trichlorophenol, 2,4,5-",
    "2,4,5-trichlorophenol":            "Trichlorophenol, 2,4,5-",
    "2.4.6-trichlorophenol":            "Trichlorophenol, 2,4,6-",
    "2,4,6-trichlorophenol":            "Trichlorophenol, 2,4,6-",
    "2.6-dichlorophenol":               "2,6-Dimethylphenol",
    "2,6-dichlorophenol":               "2,6-Dimethylphenol",
    "4.6-dinitro-2-methylphenol":       "4,6-Dinitro-o-cresol",
    "4,6-dinitro-2-methylphenol":       "4,6-Dinitro-o-cresol",
    "3-nitroaniline":                   "3,5-Dinitroaniline",
    "bis(2-chloroisopropyl)ether":      "Bis(2-chloro-1-methylethyl) ether",
    "n-nitrosodi-n-propylamine":        "N-Nitroso-di-N-propylamine",
    "di-n-butyl phthalate":             "Dibutyl Phthalate",
    "butyl benzyl phthalate":           "Butyl Benzyl Phthalate",
    "di-n-octyl phthalate":             "di-N-Octyl Phthalate",
    "diethyl phthalate":                "Diethyl Phthalate",
    "dimethyl phthalate":               "Dimethylterephthalate",
    "4-chloro-3-methylphenol":          "p-chloro-m-Cresol",
    "1.4-dioxane":                      "1,4-Dioxane",
    "1.2-dichlorobenzene":              "1,2-Dichlorobenzene",
    "1.4-dichlorobenzene":              "1,4-Dichlorobenzene",
    "1.1-dichloroethane":               "1,1-Dichloroethane",
    "1.2-dichloropropane":              "1,2-Dichloropropane",
    "2.4-dimethylphenol":               "2,4-Dimethylphenol",
    "2.4-dichlorophenol":               "2,4-Dichlorophenol",
    "2.4-dinitrophenol":                "2,4-Dinitrophenol",
    "2.4-dinitrotoluene":               "2,4-Dinitrotoluene",
    "2.6-dinitrotoluene":               "2,6-Dinitrotoluene",
    "bis(2-ethylhexyl)phthalate":       "Bis(2-ethylhexyl)phthalate",
    "1,1'-biphenyl":                    "1,1'-Biphenyl",
    "n-butylbenzene":                   "n-Butylbenzene",
    "vinyl chloride":                   "Vinyl Chloride",
}

_ALS_PFAS_ALIAS = {
    "2,3,3,3-tetrafluoro-2-(heptafluoropropoxy)propanoic acid (hfpo-da)":
        "hexafluoropropylene oxide dimer acid (hfpo-da)",
    "7h-perfluoroheptanoic acid (hpfhpa)":      "perfluoroheptanoic acid (pfhpa)",
    "perfluorobutane sulfonic acid (pfbs)":     "perfluorobutanesulfonic acid (pfbs)",
    "perfluorobutane sulfonate (pfbs)":         "perfluorobutanesulfonic acid (pfbs)",
    "perfluorohexane sulfonic acid (pfhxs)":    "perfluorohexanesulfonic acid (pfhxs)",
    "perfluorohexane sulfonate (pfhxs)":        "perfluorohexanesulfonic acid (pfhxs)",
    "perfluorooctane sulfonic acid (pfos)":     "perfluorooctanesulfonic acid (pfos)",
    "perfluorooctane sulfonate (pfos)":         "perfluorooctanesulfonic acid (pfos)",
    "perfluorooctadecanoic acid (pfocda)":      "perfluorooctadecanoic acid (pfoda)",
    "perfluoroundecanoic acid (pfunda)":        "perfluoroundecanoic acid (pfuda)",
    "perfluorotetradecanoic acid (pfcpda)":     "perfluorotetradecanoic acid (pfteta)",
    "perfluorodecane sulfonic acid (pfds)":     "perfluorodecanesulfonic acid (pfds)",
    "perfluoroheptane sulfonic acid (pfhps)":   "perfluoroheptanesulfonic acid (pfhps)",
    "perfluoropentane sulfonic acid (pfpes)":   "perfluoropentanesulfonic acid (pfpes)",
    "perfluorooctane sulfonamide (fosa)":       "perfluorooctanesulfonamide (fosa)",
    "perfluoropentanoic acid (pfpea)":          "perfluoropentanoic acid (pfpea)",
    "perfluorodecanoic acid (pfda)":            "perfluorodecanoic acid (pfda)",
    "perfluorododecanoic acid (pfdoda)":        "perfluorododecanoic acid (pfdoda)",
    "perfluoroheptanoic acid (pfhpa)":          "perfluoroheptanoic acid (pfhpa)",
    "perfluorotridecanoic acid (pftrda)":       "perfluorotridecanoic acid (pftrda)",
    "perfluorooctanesulfonic acid (pfos)":      "perfluorooctanesulfonic acid (pfos)",
}

_ALS_CANONICAL_KEY = "__CANONICAL_MAP_INTERNAL__"


# ── Utility helpers ───────────────────────────────────────────────

def _als_norm(s: str) -> str:
    s = "" if s is None else str(s).strip().lower()
    return re.sub(r"\s+", " ", s.replace("\xa0", " "))


def _als_sort_key(sid: str) -> int:
    m = re.match(r"S-?(\d+)", str(sid), re.I)
    return int(m.group(1)) if m else 9999


def _als_to_float(v) -> float | None:
    s = str(v).strip() if v is not None else ""
    try:
        return float(s.lstrip("<>").strip())
    except Exception:
        return None


def _als_check_exceed(val_str, vsl, tier1) -> str | None:
    if not val_str or str(val_str).strip().startswith("<"):
        return None
    f = _als_to_float(val_str)
    if f is None:
        return None
    try:
        t1f = float(tier1) if tier1 is not None else None
        vf  = float(vsl)   if vsl   is not None else None
        if t1f and t1f > 0 and f > t1f:
            return "tier1"
        if vf and vf > 0 and f > vf:
            return "vsl"
    except Exception:
        pass
    return None


def _als_parse_sample(sname: str):
    sname = str(sname).strip()
    if "DUP" in sname.upper():
        return None, None
    m = re.match(r"^(S\d+[A-Za-z0-9]*)\s*\(([0-9.]+)\)", sname)
    if m:
        return m.group(1), float(m.group(2))
    m = re.match(r"^(S\d+)-([0-9]+\.?[0-9]*)$", sname)
    if m:
        return m.group(1), float(m.group(2))
    m = re.match(r"^(\d+[\.\d]*)\s*\(([0-9.]+)\)", sname)
    if m:
        return m.group(1), float(m.group(2))
    return sname, None


def _als_canonical_compound(name: str) -> str:
    s = _als_norm(name)
    s = s.replace("ethylene", "ethen").replace("ethene", "ethen")
    s = re.sub(r"\s*\([^)]+\)\s*$", "", s)
    nums = re.findall(r"\d+", s)
    num_part = ",".join(nums) if nums else ""
    base_no_nums = re.sub(r"[0-9.,/\-]", " ", s)
    base_no_nums = re.sub(r"\s+", "", base_no_nums)
    return (num_part + " " + base_no_nums).strip()


# ── Threshold helpers ─────────────────────────────────────────────

def _match_thresh_simple(compound: str, thresh_dict: dict) -> dict:
    key = _als_norm(compound)
    for k in (key, key.replace(".", ","), key.replace(",", ".")):
        if k in thresh_dict and k != _ALS_CANONICAL_KEY:
            return thresh_dict[k]
    return {}


def match_threshold(compound_name: str, thresh_dict: dict) -> dict:
    key = _als_norm(compound_name)
    canon_map = thresh_dict.get(_ALS_CANONICAL_KEY)

    for k in (key, key.replace(".", ","), key.replace(",", ".")):
        if k in thresh_dict and k != _ALS_CANONICAL_KEY:
            return thresh_dict[k]

    aliased_voc = _ALS_VOC_ALIAS.get(key) or _ALS_VOC_ALIAS.get(key.replace(".", ","))
    if aliased_voc:
        a_key = _als_norm(aliased_voc)
        for k in (a_key, a_key.replace(".", ","), a_key.replace(",", ".")):
            if k in thresh_dict and k != _ALS_CANONICAL_KEY:
                return thresh_dict[k]

    aliased = _ALS_PFAS_ALIAS.get(key)
    if aliased:
        a_key = _als_norm(aliased)
        for k in (a_key, a_key.replace(".", ","), a_key.replace(",", ".")):
            if k in thresh_dict and k != _ALS_CANONICAL_KEY:
                return thresh_dict[k]

    stripped = re.sub(r"\s*\([A-Z0-9:_\-]+\)\s*$", "", compound_name).strip().lower()
    for k, v in thresh_dict.items():
        if k == _ALS_CANONICAL_KEY:
            continue
        k_s = re.sub(r"\s*\([A-Z0-9:_\-]+\)\s*$", "", k).strip().lower()
        if len(stripped) > 8 and stripped == k_s:
            return v

    if canon_map:
        ck = _als_canonical_compound(compound_name)
        hit = canon_map.get(ck)
        if hit:
            return hit

    return {}


def load_threshold_file(file_bytes: bytes) -> dict:
    wb = load_workbook(io.BytesIO(file_bytes), data_only=True)
    ws = wb.active
    thresh: dict = {}
    for row in ws.iter_rows(min_row=6, values_only=True):
        if not row[0]:
            continue
        name = str(row[0]).strip()
        cas  = str(row[1]).strip() if row[1] else "-"

        def _g(ci):
            return (row[ci] if ci < len(row) and row[ci] is not None
                    and str(row[ci]) not in ("NA", "") else None)

        thresh[_als_norm(name)] = {
            "name": name, "cas": cas,
            "units": str(row[3]) if row[3] else "mg/kg",
            "VSL": _g(4),
            "Ind_A_06": _g(8), "Ind_A_6p": _g(9),
            "Ind_B":    _g(10), "Res_A_06": _g(11),
            "Res_A_6p": _g(12), "Res_B":    _g(13),
        }

    canon_map: dict = {}
    for k, v in thresh.items():
        ck = _als_canonical_compound(k)
        canon_map.setdefault(ck, v)
    thresh[_ALS_CANONICAL_KEY] = canon_map
    return thresh


def get_tier1_col(land_use: str, aquifer: str, depth: str) -> str:
    ind = "industrial" in land_use.lower()
    b   = "b-1" in aquifer.lower()
    if b:
        return "Ind_B" if ind else "Res_B"
    deep = ">6" in depth
    if ind:
        return "Ind_A_06" if not deep else "Ind_A_6p"
    return "Res_A_06" if not deep else "Res_A_6p"


def tier1_label(land_use: str, aquifer: str, depth: str) -> str:
    return f"TIER 1\n{land_use}\n{aquifer}\n{depth}"


# ── Low-level Word helpers for the ALS builder ────────────────────

def _als_set_cell_borders(cell, color: str = "000000", size: int = 4) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _als_cell_text(
    cell,
    text,
    bold: bool = False,
    size: int = 8,
    color: str | None = None,
    align: str = "center",
    rtl: bool = False,
    bg: str | None = None,
) -> None:
    from docx.enum.table import WD_ALIGN_VERTICAL as _VAL
    cell.vertical_alignment = _VAL.CENTER
    if bg:
        _set_cell_bg(cell, bg)
    _als_set_cell_borders(cell)

    p = cell.paragraphs[0]
    p.clear()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if rtl:
        pPr = p._p.get_or_add_pPr()
        bidi = OxmlElement("w:bidi")
        pPr.append(bidi)

    run = p.add_run(str(text) if text is not None else "")
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _als_set_section_props(section, page_size: str, landscape: bool) -> None:
    w_twips, h_twips = _ALS_PAGE_SIZES[page_size]
    if landscape:
        section.page_width  = Twips(h_twips)
        section.page_height = Twips(w_twips)
        section.orientation = WD_ORIENT.LANDSCAPE
    else:
        section.page_width  = Twips(w_twips)
        section.page_height = Twips(h_twips)
        section.orientation = WD_ORIENT.PORTRAIT
    m = _ALS_MARGIN
    section.top_margin    = Twips(m)
    section.bottom_margin = Twips(m)
    section.left_margin   = Twips(m)
    section.right_margin  = Twips(m)


def _als_content_width(page_size: str, landscape: bool) -> int:
    w_twips, h_twips = _ALS_PAGE_SIZES[page_size]
    base = h_twips if landscape else w_twips
    return base - 2 * _ALS_MARGIN


def _als_add_section_break(doc: Document, page_size: str, landscape: bool):
    new_sec = doc.add_section()
    _als_set_section_props(new_sec, page_size, landscape)
    return new_sec


def _als_add_title(doc: Document, title: str, part_str: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(f"{title}   {part_str}")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11)


def _als_add_legend(doc: Document, has_yellow: bool, has_orange: bool) -> None:
    if not has_yellow and not has_orange:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(0)

    if has_yellow:
        r1 = p.add_run("■ ")
        r1.font.name = "Arial"
        r1.font.size = Pt(8)
        r1.font.color.rgb = RGBColor.from_string(_ALS_COLOR_YELLOW)
        r2 = p.add_run("בצהוב - חריגה מערך הסף VSL    ")
        r2.font.name = "Arial"
        r2.font.size = Pt(8)

    if has_orange:
        r3 = p.add_run("■ ")
        r3.font.name = "Arial"
        r3.font.size = Pt(8)
        r3.font.color.rgb = RGBColor.from_string(_ALS_COLOR_ORANGE)
        r4 = p.add_run("בכתום - חריגה מערך הסף TIER 1")
        r4.font.name = "Arial"
        r4.font.size = Pt(8)


def _als_set_table_width(table, width_twips: int) -> None:
    tbl  = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),    str(width_twips))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


def _als_set_col_width(cell, width_twips: int) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcW")):
        tcPr.remove(old)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(width_twips))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


# ── Table-data builders ───────────────────────────────────────────

def _build_tph_table_data(df: pd.DataFrame, thresh_dict: dict, t1col: str, t1lbl: str):
    def _is_dro(c):
        if "(dro)" in c or "- dro" in c or c.strip() == "dro":
            return True
        if "(oro)" in c or "- oro" in c:
            return False
        return "c10" in c and "c28" in c and "c40" not in c

    def _is_oro(c):
        if "(oro)" in c or "- oro" in c or c.strip() == "oro":
            return True
        if "(dro)" in c or "- dro" in c:
            return False
        return ("c24" in c and "c40" in c) or ("c28" in c and "c40" in c)

    def _is_total(c):
        if any(x in c for x in ["(dro)", "(oro)", "- dro", "- oro"]):
            return False
        if "c10" in c and "c40" in c:
            return True
        return "total" in c and ("tph" in c or "hydrocarbon" in c)

    def _get(compound):
        t = _match_thresh_simple(compound, thresh_dict)
        return t.get("VSL"), t.get(t1col), t.get("cas", "-")

    vsl_d, t1_d, _ = _get("C10 - C28 Fraction (DRO)")
    vsl_o, t1_o, _ = _get("C24 - C40 Fraction (ORO)")
    vsl_t, t1_t, _ = _get("TPH - DRO + ORO (Tier 1)")
    vv = [v for v in [vsl_d, vsl_o, vsl_t] if v]
    tt = [v for v in [t1_d, t1_o, t1_t] if v]
    vsl_tot = min(float(v) for v in vv) if vv else 350
    t1_tot  = min(float(v) for v in tt) if tt else 350

    headers = [
        ["שם קידוח", "עומק", "TPH DRO", "TPH ORO", "Total TPH"],
        ["", "יחידות", "mg/kg", "mg/kg", "mg/kg"],
        ["", "CAS", "C10-C40", "C10-C40", "C10-C40"],
        ["", "VSL", str(vsl_tot), str(vsl_tot), str(vsl_tot)],
        ["", t1lbl.replace("\n", " "), str(t1_tot), str(t1_tot), str(t1_tot)],
    ]

    pivoted: dict = {}
    for _, r in df.iterrows():
        k = (r["sample_id"], r["depth"])
        if k not in pivoted:
            pivoted[k] = {"DRO": "", "ORO": "", "TOT": "", "DRO_f": None,
                          "ORO_f": None, "DRO_lor": None, "ORO_lor": None}
        c = r["compound_lower"]
        if _is_dro(c) and not pivoted[k]["DRO"]:
            pivoted[k]["DRO"]     = r["result_str"]
            pivoted[k]["DRO_f"]   = r["result"]
            pivoted[k]["DRO_lor"] = r.get("lor_val")
        elif _is_oro(c) and not pivoted[k]["ORO"]:
            pivoted[k]["ORO"]     = r["result_str"]
            pivoted[k]["ORO_f"]   = r["result"]
            pivoted[k]["ORO_lor"] = r.get("lor_val")
        elif _is_total(c) and not pivoted[k]["TOT"]:
            pivoted[k]["TOT"] = r["result_str"]

    rows_out = []
    prev_sid = None
    for (sid, depth_val), v in sorted(pivoted.items(),
                                       key=lambda x: (_als_sort_key(x[0][0]), x[0][1] or 0)):
        if v["TOT"]:
            total_s = v["TOT"]
        else:
            dro_lor   = v["DRO"] and str(v["DRO"]).startswith("<")
            oro_lor   = v["ORO"] and str(v["ORO"]).startswith("<")
            dro_empty = not v["DRO"]
            oro_empty = not v["ORO"]
            dro_num = v["DRO_lor"] if dro_lor and v["DRO_lor"] is not None else (v["DRO_f"] or 0)
            oro_num = v["ORO_lor"] if oro_lor and v["ORO_lor"] is not None else (v["ORO_f"] or 0)
            total_f = dro_num + oro_num
            if (dro_lor or dro_empty) and (oro_lor or oro_empty) and not (dro_empty and oro_empty):
                total_s = f"<{total_f:.0f}"
            else:
                total_s = f"{total_f:.0f}"

        sid_display = sid if sid != prev_sid else ""
        prev_sid    = sid
        rows_out.append({
            "values": [sid_display, str(depth_val) if depth_val else "",
                       v["DRO"], v["ORO"], total_s],
            "colors": [None, None,
                       _als_check_exceed(v["DRO"],   vsl_tot, t1_tot),
                       _als_check_exceed(v["ORO"],   vsl_tot, t1_tot),
                       _als_check_exceed(total_s,    vsl_tot, t1_tot)],
        })

    return headers, rows_out


def _build_metals_table_data(df: pd.DataFrame, thresh_dict: dict, t1col: str, t1lbl: str):
    df = df.copy()
    df["sym"] = df["compound_lower"].map(_ALS_METAL_MAP)
    df = df[df["sym"].notna()]
    if df.empty:
        return [["אין נתוני מתכות"]], []

    present = set(df["sym"].unique())
    metals  = [m for m in _ALS_METALS_ORDER if m in present] + sorted(present - set(_ALS_METALS_ORDER))

    mt: dict = {}
    for key, v in thresh_dict.items():
        if key == _ALS_CANONICAL_KEY:
            continue
        sym = _ALS_THRESH_METAL_MAP.get(key)
        if sym and sym not in mt:
            mt[sym] = {"vsl": v.get("VSL"), "tier1": v.get(t1col), "cas": v.get("cas", "-")}

    headers = [
        ["שם קידוח", "עומק"] + metals,
        ["", "יחידות"] + ["mg/kg"] * len(metals),
        ["", "CAS"]     + [mt.get(m, {}).get("cas",   "-") for m in metals],
        ["", "VSL"]     + [str(mt.get(m, {}).get("vsl",   "-")) for m in metals],
        ["", t1lbl.replace("\n", " ")] + [str(mt.get(m, {}).get("tier1", "-")) for m in metals],
    ]

    pt = df.pivot_table(index=["sample_id", "depth"], columns="sym",
                        values="result_str", aggfunc="first")
    pt = pt.reindex(sorted(pt.index, key=lambda x: (_als_sort_key(x[0]), x[1] or 0)))

    rows_out = []
    prev_sid = None
    for (sid, depth_val), row_data in pt.iterrows():
        sid_display = sid if sid != prev_sid else ""
        prev_sid = sid
        values = [sid_display, str(depth_val) if depth_val else ""]
        colors: list = [None, None]
        for sym in metals:
            val = row_data.get(sym, "") or ""
            val = "" if str(val) == "nan" else str(val)
            colors.append(_als_check_exceed(val, mt.get(sym, {}).get("vsl"),
                                            mt.get(sym, {}).get("tier1")))
            values.append(val)
        rows_out.append({"values": values, "colors": colors})

    return headers, rows_out


def _build_generic_table_data(df: pd.DataFrame, thresh_dict: dict,
                              t1col: str, t1lbl: str):
    pairs = sorted(df[["sample_id", "depth"]].drop_duplicates().values.tolist(),
                   key=lambda x: (_als_sort_key(x[0]), -(x[1] or 0)))
    compounds = df["compound"].unique().tolist()

    headers = [
        ["שם התרכובת", "CAS", "VSL", t1lbl.replace("\n", " "), "יחידות"]
        + [f"{sid}\n{d}" for sid, d in pairs],
    ]

    als_data: dict = {}
    for _, r in df.iterrows():
        k = str(r["compound"]).strip()
        als_data.setdefault(k, {})[(r["sample_id"], r["depth"])] = r["result_str"]

    rows_out = []
    for cmp in compounds:
        t     = _match_thresh_simple(cmp, thresh_dict)
        vsl   = t.get("VSL")
        tier1 = t.get(t1col)
        cas   = t.get("cas", "-")
        sub   = df[df["compound"] == cmp]
        unit  = sub["unit"].iloc[0] if len(sub) > 0 else "mg/kg"
        values = [cmp, str(cas),
                  str(vsl) if vsl else "-",
                  str(tier1) if tier1 else "-",
                  str(unit)]
        colors: list = [None, None, None, None, None]
        for sid, depth_val in pairs:
            rs = als_data.get(cmp, {}).get((sid, depth_val), "")
            colors.append(_als_check_exceed(rs, vsl, tier1))
            values.append(str(rs) if rs else "")
        rows_out.append({"values": values, "colors": colors})

    return headers, rows_out


# ── Core table-to-document writer ────────────────────────────────

def _add_table_to_doc(
    doc: Document,
    headers: list,
    data_rows: list,
    title: str,
    page_size: str,
    landscape: bool,
    is_first_section: bool = False,
) -> None:
    if not data_rows:
        return

    n_cols    = len(headers[0])
    content_w = _als_content_width(page_size, landscape)

    if n_cols <= 5:
        col_w_first = int(content_w * 0.25)
        col_w_rest  = int((content_w - col_w_first) / max(n_cols - 1, 1))
        col_widths  = [col_w_first] + [col_w_rest] * (n_cols - 1)
    else:
        col_w_first  = int(content_w * 0.18)
        col_w_second = int(content_w * 0.08)
        remaining    = content_w - col_w_first - col_w_second
        col_w_rest   = int(remaining / max(n_cols - 2, 1))
        col_widths   = [col_w_first, col_w_second] + [col_w_rest] * (n_cols - 2)

    total = sum(col_widths)
    if total != content_w:
        col_widths[-1] += content_w - total

    page_h   = _ALS_PAGE_SIZES[page_size][1] if not landscape else _ALS_PAGE_SIZES[page_size][0]
    avail_h  = page_h - 2 * _ALS_MARGIN
    n_hdrs   = len(headers)
    rpp      = max(5, int((avail_h - 400 - 300 - n_hdrs * 600) / 280))

    chunks      = [data_rows[i:i + rpp] for i in range(0, len(data_rows), rpp)]
    total_parts = len(chunks)

    for part_idx, chunk in enumerate(chunks):
        part_str = (f"(חלק {part_idx + 1} מתוך {total_parts})"
                    if total_parts > 1 else "")

        if not (is_first_section and part_idx == 0):
            _als_add_section_break(doc, page_size, landscape)

        _als_add_title(doc, title, part_str)

        n_rows_total = n_hdrs + len(chunk)
        table = doc.add_table(rows=n_rows_total, cols=n_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _als_set_table_width(table, content_w)

        for hi, hrow in enumerate(headers):
            for ci, val in enumerate(hrow):
                cell = table.cell(hi, ci)
                _als_set_col_width(cell, col_widths[ci])
                bg = _ALS_COLOR_HEADER2 if hi == 0 else _ALS_COLOR_HEADER
                _als_cell_text(cell, val, bold=True, size=8, bg=bg, rtl=True)

        has_yellow = False
        has_orange = False
        for ri, row_data in enumerate(chunk):
            tr_idx = n_hdrs + ri
            for ci, (val, color) in enumerate(zip(row_data["values"], row_data["colors"])):
                cell = table.cell(tr_idx, ci)
                _als_set_col_width(cell, col_widths[ci])
                bg = None
                if color == "vsl":
                    bg = _ALS_COLOR_YELLOW
                    has_yellow = True
                elif color == "tier1":
                    bg = _ALS_COLOR_ORANGE
                    has_orange = True
                _als_cell_text(cell, val, bold=bool(color), size=8,
                               bg=bg or _ALS_COLOR_WHITE,
                               align="right" if ci == 0 else "center",
                               rtl=True)

        _als_add_legend(doc, has_yellow, has_orange)


# ── Public builder ────────────────────────────────────────────────

def build_word_report(
    table_configs: list[dict],
    thresh_dict: dict,
    t1col: str,
    t1lbl: str,
) -> bytes:
    """
    Build a Word document from parsed ALS data.

    Parameters
    ----------
    table_configs : list of dicts, each with:
        type      – "TPH" | "Metals" | "VOC+SVOC" | "PFAS"
        df        – pd.DataFrame from parse_als_file
        title     – str, Hebrew table title
        page_size – "A4" | "Tabloid"
        landscape – bool
    thresh_dict : dict from load_threshold_file
    t1col       : str from get_tier1_col
    t1lbl       : str from tier1_label
    Returns
    -------
    bytes of the .docx file
    """
    doc = Document()

    for section in doc.sections:
        section.page_width  = Twips(_ALS_PAGE_SIZES["A4"][0])
        section.page_height = Twips(_ALS_PAGE_SIZES["A4"][1])
        for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
            setattr(section, attr, Twips(_ALS_MARGIN))

    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    is_first = True

    for cfg in table_configs:
        df        = cfg["df"]
        ttype     = cfg["type"]
        title     = cfg["title"]
        page_size = cfg.get("page_size", "A4")
        landscape = cfg.get("landscape", False)

        if df is None or df.empty:
            continue

        if ttype == "TPH":
            headers, data_rows = _build_tph_table_data(df, thresh_dict, t1col, t1lbl)
        elif ttype == "Metals":
            headers, data_rows = _build_metals_table_data(df, thresh_dict, t1col, t1lbl)
        else:
            headers, data_rows = _build_generic_table_data(df, thresh_dict, t1col, t1lbl)

        if not data_rows:
            continue

        if is_first:
            _als_set_section_props(doc.sections[0], page_size, landscape)
            is_first = False

        _add_table_to_doc(doc, headers, data_rows, title, page_size, landscape,
                          is_first_section=(len(doc.element.body) <= 2))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ── ALS file parser ───────────────────────────────────────────────

def parse_als_file(file_bytes: bytes, filename: str):
    """
    Parse an ALS Excel report into a flat DataFrame.

    Returns (DataFrame, None) on success, (None, error_str) on failure.

    DataFrame columns:
        sample_id, depth, compound, compound_lower, unit,
        lor, result, result_str, lor_val, group, source
    """
    try:
        wb = load_workbook(io.BytesIO(file_bytes), data_only=True)
    except Exception as e:
        return None, str(e)

    main = next(
        (wb[n] for n in wb.sheetnames if "Client" in n and "SOIL" in n),
        wb.worksheets[0],
    )
    rows = list(main.iter_rows(values_only=True))

    sid_row_idx = next(
        (i for i, r in enumerate(rows)
         if any("Client Sample ID" in str(v) for v in r if v)),
        None,
    )
    if sid_row_idx is None:
        return None, "לא נמצאה שורת Client Sample ID"

    sid_row = rows[sid_row_idx]
    sid_label_col = next(
        ci for ci, v in enumerate(sid_row) if v and "Client Sample ID" in str(v)
    )
    col2sample = {
        ci: str(v).strip()
        for ci, v in enumerate(sid_row)
        if ci > sid_label_col and v and str(v).strip() not in ("", "None")
    }

    ph_idx = next(
        (i for i, r in enumerate(rows) if r and r[0] == "Parameter"),
        None,
    )
    if ph_idx is None:
        return None, "לא נמצאה שורת Parameter"

    param_row = rows[ph_idx]
    unit_col = next(
        (ci for ci, v in enumerate(param_row)
         if v and str(v).strip().lower() == "unit"),
        2,
    )
    lor_col = next(
        (ci for ci, v in enumerate(param_row)
         if v and str(v).strip().lower() == "lor"),
        unit_col + 1,
    )

    records = []
    group   = "Unknown"

    for row in rows[ph_idx + 1:]:
        p = row[0] if row else None
        if not p or str(p).strip() in ("", "None"):
            continue

        method = row[1] if len(row) > 1 else None
        if not method or str(method).strip() in ("", "None"):
            group = str(p).strip()
            continue

        u   = row[unit_col] if unit_col < len(row) else None
        lor = row[lor_col]  if lor_col  < len(row) else None

        for ci, sname in col2sample.items():
            sid, depth_val = _als_parse_sample(sname)
            if sid is None:
                continue
            val = row[ci] if ci < len(row) else None
            rs  = str(val).strip() if val is not None else ""
            if rs in ("", "None"):
                continue

            result  = None
            lor_val = None
            if rs.startswith("<"):
                result = 0.0
                try:
                    lor_val = float(rs[1:].strip())
                except Exception:
                    lor_val = 0.0
            else:
                try:
                    result = float(rs)
                except Exception:
                    result = None

            if result is not None:
                records.append({
                    "sample_id":      sid,
                    "depth":          depth_val,
                    "compound":       str(p).strip(),
                    "compound_lower": _als_norm(p),
                    "unit":           str(u).strip() if u else "mg/kg",
                    "lor":            lor,
                    "result":         result,
                    "result_str":     rs,
                    "lor_val":        lor_val,
                    "group":          group,
                    "source":         filename,
                })

    if not records:
        return None, "לא נמצאו נתונים"
    return pd.DataFrame(records), None
