"""
core/annual_gw_report.py
------------------------
Annual groundwater monitoring report generator.

Workflow:
1. Parse the most recent periodic report (Word .docx) to extract:
   - Historical cumulative tables (field + lab results per borehole)
   - Site metadata (name, station number, boreholes list)
2. Parse new lab reports (PDF/Excel, any supported lab) for the current year
3. Merge new results into the historical tables
4. Generate updated Word report based on the periodic report structure

The output Word document updates:
  - Cover page year and report type ("שנתי")  
  - Table of field measurements (all boreholes, all dates including new)
  - Table of lab results BTEX/MTBE (all boreholes, all dates including new)
  - Summary section with updated latest sampling date and findings
"""

from __future__ import annotations

import io
import re
from datetime import datetime
from typing import Any

# ── Data structures ──────────────────────────────────────────────────────────

class SamplingEvent:
    """Single sampling event for one borehole."""
    def __init__(self, borehole: str, date: str):
        self.borehole  = borehole   # e.g. "מת-2"
        self.date      = date       # e.g. "03.08.25"
        # Field params
        self.ph:          str | None = None
        self.ec:          str | None = None
        self.temp:        str | None = None
        self.do:          str | None = None
        self.turbidity:   str | None = None
        self.redox:       str | None = None
        self.water_level: str | None = None
        self.total_depth: str | None = None
        self.sample_depth:str | None = None
        self.method:      str | None = None  # "Low flow" / "ביילר" / note
        # Lab results (mg/L)
        self.benzene:      str | None = None
        self.toluene:      str | None = None
        self.ethylbenzene: str | None = None
        self.xylene:       str | None = None
        self.mtbe:         str | None = None
        self.naphthalene:  str | None = None
        # Source
        self.source: str = ""  # "word" / lab name

    def date_obj(self) -> datetime | None:
        for fmt in ("%d.%m.%y", "%d.%m.%Y", "%d/%m/%y", "%d/%m/%Y"):
            try:
                return datetime.strptime(self.date.strip(), fmt)
            except ValueError:
                continue
        return None

    def __repr__(self):
        return f"<SamplingEvent {self.borehole} {self.date}>"


class GWReportData:
    """All data extracted from a groundwater report."""
    def __init__(self):
        self.site_name:     str = ""
        self.station_no:    str = ""
        self.boreholes:     list[str] = []       # ordered list of borehole names
        self.events:        list[SamplingEvent] = []
        self.target_values: dict[str, str] = {}  # compound → target value string
        self.report_type:   str = ""             # "שנתי" / "תקופתי"
        self.report_year:   str = ""
        self.report_date:   str = ""
        self.author:        str = ""
        self.approver:      str = ""
        # Raw Word XML for template editing
        self.word_xml_dir:  str | None = None
        # Metadata text blocks for summary
        self.background_text: str = ""

    def get_event(self, borehole: str, date: str) -> SamplingEvent | None:
        for e in self.events:
            if e.borehole == borehole and e.date.strip() == date.strip():
                return e
        return None

    def get_or_create_event(self, borehole: str, date: str) -> SamplingEvent:
        ev = self.get_event(borehole, date)
        if ev is None:
            ev = SamplingEvent(borehole, date)
            self.events.append(ev)
            if borehole not in self.boreholes:
                self.boreholes.append(borehole)
        return ev

    def events_for_borehole(self, borehole: str) -> list[SamplingEvent]:
        evs = [e for e in self.events if e.borehole == borehole]
        evs.sort(key=lambda e: e.date_obj() or datetime.min)
        return evs

    def all_dates_sorted(self) -> list[str]:
        seen = {}
        for e in self.events:
            d = e.date_obj()
            if d:
                seen[e.date.strip()] = d
        return [k for k, _ in sorted(seen.items(), key=lambda x: x[1])]

    def new_dates(self, cutoff_date: str) -> list[str]:
        """Return dates strictly after cutoff_date."""
        cutoff = None
        for fmt in ("%d.%m.%y", "%d.%m.%Y"):
            try:
                cutoff = datetime.strptime(cutoff_date.strip(), fmt)
                break
            except ValueError:
                pass
        if cutoff is None:
            return []
        return [d for d in self.all_dates_sorted()
                if (dt := self._parse_date(d)) and dt > cutoff]

    @staticmethod
    def _parse_date(s: str) -> datetime | None:
        for fmt in ("%d.%m.%y", "%d.%m.%Y"):
            try:
                return datetime.strptime(s.strip(), fmt)
            except ValueError:
                pass
        return None


# ── Word document parser ─────────────────────────────────────────────────────

def parse_word_report(docx_bytes: bytes, docx_name: str = "") -> GWReportData:
    """
    Parse a periodic or annual groundwater Word report.
    Returns GWReportData with historical tables populated.
    """
    try:
        import docx as _docx
    except ImportError:
        import subprocess, sys
        subprocess.check_call([sys.executable, "-m", "pip", "install",
                               "python-docx", "--break-system-packages", "-q"])
        import docx as _docx

    data = GWReportData()
    doc  = _docx.Document(io.BytesIO(docx_bytes))

    # ── Extract site metadata from paragraphs ────────────────────────────────
    full_text = "\n".join(p.text for p in doc.paragraphs)

    # Report type
    if "שנתי" in full_text[:500]:
        data.report_type = "שנתי"
    else:
        data.report_type = "תקופתי"

    # Year
    year_m = re.search(r'\b(202\d)\b', full_text[:300])
    if year_m:
        data.report_year = year_m.group(1)

    # Month/year from title area
    month_m = re.search(r'(ינואר|פברואר|מרץ|אפריל|מאי|יוני|יולי|אוגוסט|ספטמבר|אוקטובר|נובמבר|דצמבר)\s+(20\d\d)', full_text[:400])
    if month_m:
        data.report_date = f"{month_m.group(1)} {month_m.group(2)}"

    # Station name
    name_m = re.search(r'תחנת\s+([\w\s"\']+?)[\n,]', full_text[:500])
    if name_m:
        data.site_name = name_m.group(1).strip()
    station_m = re.search(r'תחנה\s+מספר\s+(\d+)', full_text[:500])
    if station_m:
        data.station_no = station_m.group(1)

    # Author / approver from cover table
    for table in doc.tables[:3]:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            text  = " ".join(cells)
            if "מחבר" in text:
                for c in row.cells:
                    t = c.text.strip()
                    if t and "מחבר" not in t and "חתימה" not in t:
                        data.author = t
            if "מאשר" in text or "בודק" in text:
                for c in row.cells:
                    t = c.text.strip()
                    if t and "מאשר" not in t and "בודק" not in t and "חתימה" not in t:
                        data.approver = t

    # ── Find target values ───────────────────────────────────────────────────
    tv_m = re.search(r'ערך\s+יעד[^|]*\|[^|]*\|([^|]+)\|([^|]+)\|([^|]+)\|([^|]+)\|([^|]+)', full_text)
    if tv_m:
        params = ["Benzene", "Toluene", "Ethyl Benzene", "Xylene", "MTBE"]
        for i, p in enumerate(params, 1):
            data.target_values[p] = tv_m.group(i).strip()
    else:
        # Defaults from drinking water standard
        data.target_values = {
            "Benzene": "0.005", "Toluene": "0.70",
            "Ethyl Benzene": "0.30", "Xylene": "0.50", "MTBE": "0.04"
        }

    # ── Parse tables ─────────────────────────────────────────────────────────
    for table in doc.tables:
        header_row = table.rows[0].cells[0].text.strip() if table.rows else ""
        _try_parse_lab_table(table, data)
        _try_parse_field_table(table, data)

    # Collect background text (first few paragraphs after title)
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    data.background_text = "\n".join(paras[:5])

    return data


def _try_parse_lab_table(table, data: GWReportData):
    """Try to parse lab results table (BTEX/MTBE) from a Word table."""
    try:
        rows = table.rows
        if len(rows) < 3:
            return

        # Look for table with BTEX columns
        header_text = " ".join(c.text for c in rows[0].cells)
        if not any(k in header_text for k in ["בנזן", "Benzene", "MTBE", "BTEX"]):
            # Check row 1
            header_text = " ".join(c.text for c in rows[1].cells)
            if not any(k in header_text for k in ["בנזן", "Benzene", "MTBE", "BTEX"]):
                return

        # Find column indices
        col_map = {}
        for ri in range(min(3, len(rows))):
            cells = [c.text.strip() for c in rows[ri].cells]
            for ci, cell in enumerate(cells):
                c_low = cell.lower()
                if "בנזן" in c_low or "benzene" in c_low:
                    col_map["benzene"] = ci
                if "טולואן" in c_low or "toluene" in c_low:
                    col_map["toluene"] = ci
                if "אתיל בנזן" in c_low or "ethyl" in c_low:
                    col_map["ethylbenzene"] = ci
                if "כסילן" in c_low or "xylene" in c_low:
                    col_map["xylene"] = ci
                if "mtbe" in c_low:
                    col_map["mtbe"] = ci
                if "נפטלן" in c_low or "naphthalene" in c_low:
                    col_map["naphthalene"] = ci
                if ("שם קידוח" in cell or "קידוח" in cell) and "שם" in cell:
                    col_map["borehole"] = ci
                if "תאריך" in cell:
                    col_map["date"] = ci
                if "מפלס" in cell or "עומק פני המים" in cell:
                    col_map["water_level"] = ci

        if not col_map.get("benzene") and not col_map.get("mtbe"):
            return

        current_borehole = ""
        for row in rows[2:]:
            cells = [c.text.strip() for c in row.cells]
            if not any(cells):
                continue

            bh_idx  = col_map.get("borehole", 0)
            dt_idx  = col_map.get("date", 1)
            bh_text = cells[bh_idx] if bh_idx < len(cells) else ""
            dt_text = cells[dt_idx] if dt_idx < len(cells) else ""

            if bh_text and re.search(r'מת|MW|מד|קידוח', bh_text, re.I):
                current_borehole = bh_text.strip().strip("*").strip()

            if not current_borehole:
                continue

            # Parse date
            date_m = re.search(r'\d{1,2}[./]\d{1,2}[./]\d{2,4}', dt_text)
            if not date_m:
                continue
            date_str = _normalize_date(date_m.group())

            ev = data.get_or_create_event(current_borehole, date_str)
            ev.source = "word"

            def _get(col_key):
                idx = col_map.get(col_key)
                if idx is not None and idx < len(cells):
                    v = cells[idx].strip()
                    return v if v else None
                return None

            ev.benzene      = _get("benzene")
            ev.toluene      = _get("toluene")
            ev.ethylbenzene = _get("ethylbenzene")
            ev.xylene       = _get("xylene")
            ev.mtbe         = _get("mtbe")
            ev.naphthalene  = _get("naphthalene")
            if col_map.get("water_level"):
                ev.water_level = _get("water_level")

    except Exception:
        pass


def _try_parse_field_table(table, data: GWReportData):
    """Try to parse field measurements table from a Word table."""
    try:
        rows = table.rows
        if len(rows) < 3:
            return

        header_text = " ".join(c.text for c in rows[0].cells) + " " + \
                      " ".join(c.text for c in rows[1].cells)
        if not any(k in header_text for k in ["pH", "EC", "עכירות", "רדוקס", "מוליכות"]):
            return

        col_map = {}
        for ri in range(min(4, len(rows))):
            cells = [c.text.strip() for c in rows[ri].cells]
            for ci, cell in enumerate(cells):
                if "ph" in cell.lower() and "units" not in cell.lower() and ci not in col_map.values():
                    col_map["ph"] = ci
                if any(k in cell for k in ["EC", "מוליכות"]):
                    col_map["ec"] = ci
                if any(k in cell for k in ["טמפרטורה", "Temp", "°C"]):
                    col_map["temp"] = ci
                if any(k in cell for k in ["חמצן מומס", "DO", "מומס"]):
                    col_map["do"] = ci
                if "עכירות" in cell or "Turb" in cell.lower():
                    col_map["turbidity"] = ci
                if "רדוקס" in cell or "Redox" in cell.lower() or "ORP" in cell:
                    col_map["redox"] = ci
                if "עומק פני המים" in cell:
                    col_map["water_level"] = ci
                if "עומק כללי" in cell or "עומק הקידוח" in cell:
                    col_map["total_depth"] = ci
                if "עומק דגימה" in cell:
                    col_map["sample_depth"] = ci
                if ("שם קידוח" in cell or ("קידוח" in cell and "שם" in cell)):
                    col_map["borehole"] = ci
                if "תאריך" in cell:
                    col_map["date"] = ci

        if not col_map.get("water_level") and not col_map.get("ph"):
            return

        current_borehole = ""
        for row in rows:
            cells = [c.text.strip() for c in row.cells]
            if not any(cells):
                continue
            bh_idx  = col_map.get("borehole", 0)
            dt_idx  = col_map.get("date", 1)
            bh_text = cells[bh_idx] if bh_idx < len(cells) else ""
            dt_text = cells[dt_idx] if dt_idx < len(cells) else ""

            if bh_text and re.search(r'מת|MW|מד|קידוח', bh_text, re.I):
                current_borehole = bh_text.strip().strip("*").strip()

            if not current_borehole:
                continue

            date_m = re.search(r'\d{1,2}[./]\d{1,2}[./]\d{2,4}', dt_text)
            if not date_m:
                continue
            date_str = _normalize_date(date_m.group())

            ev = data.get_or_create_event(current_borehole, date_str)

            def _get(col_key):
                idx = col_map.get(col_key)
                if idx is not None and idx < len(cells):
                    v = cells[idx].strip()
                    return v if v else None
                return None

            if _get("ph"):        ev.ph           = _get("ph")
            if _get("ec"):        ev.ec           = _get("ec")
            if _get("temp"):      ev.temp         = _get("temp")
            if _get("do"):        ev.do           = _get("do")
            if _get("turbidity"): ev.turbidity    = _get("turbidity")
            if _get("redox"):     ev.redox        = _get("redox")
            if _get("water_level"): ev.water_level  = _get("water_level")
            if _get("total_depth"):  ev.total_depth  = _get("total_depth")
            if _get("sample_depth"): ev.sample_depth = _get("sample_depth")

    except Exception:
        pass


# ── Lab report parser (PDF/Excel → SamplingEvents) ──────────────────────────

def parse_lab_reports(files: list[tuple[str, bytes]]) -> list[SamplingEvent]:
    """
    Parse one or more lab reports (PDF or Excel) from any supported lab.
    Returns list of SamplingEvents with lab results populated.
    """
    import sys, os
    sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))

    try:
        from parsers import get_parser, auto_detect_lab, auto_detect_category
    except ImportError:
        return []

    events = []
    for fname, fdata in files:
        try:
            lab = auto_detect_lab(fname, fdata)
            cat = auto_detect_category(fname, fdata, lab) or "groundwater"
            if not lab:
                # Try groundwater labs manually
                if b"bactochem" in fdata[:2000].lower() or b"\xd1\xd0\xd8\xd8\xd5\xd9\xd0\xd9" in fdata[:2000]:
                    lab = "bactochem"
                elif b"aminolab" in fdata[:2000].lower() or b'\xe0\xee\xe9\xf0\xe5' in fdata[:2000]:
                    lab = "aminolab"

            parser  = get_parser(lab, cat)
            records = parser.parse(io.BytesIO(fdata))

            # Build sample_id → borehole_name map (for Bactochem PDFs)
            sid_to_bh   = _build_sample_borehole_map(fdata, fname)
            # Build sample_id → sampling_date map from PDF
            pdf_date    = _extract_sampling_date_from_pdf(fdata) if fname.lower().endswith(".pdf") else ""

            # Group records by (borehole, date)
            grouped: dict[tuple, SamplingEvent] = {}
            for rec in records:
                if rec.get("analysis_type") not in ("GW_VOC", "GW_FIELD_PARAMS",
                                                      "GW_BTEX", "GW_MTBE"):
                    if rec.get("analysis_type") and "GW" not in str(rec.get("analysis_type")):
                        continue

                bh   = _extract_borehole_from_rec(rec)
                # Override with PDF-extracted borehole map if available
                sid_raw = rec.get('sample_id', '')
                if sid_raw in sid_to_bh:
                    bh = sid_to_bh[sid_raw]
                date = _extract_date_from_sid_or_rec(rec)
                # Use PDF-level sampling date if individual record has no date
                if not __import__('re').search(r'\d{1,2}[./]\d{1,2}[./]\d{2,4}', date) and pdf_date:
                    date = pdf_date
                key  = (bh, date)

                if key not in grouped:
                    ev = SamplingEvent(bh, date)
                    ev.source = lab or fname
                    grouped[key] = ev

                ev   = grouped[key]
                cmp  = (rec.get("compound") or "").strip()
                val  = rec.get("value")
                flag = rec.get("flag", "")
                disp = _format_val(val, flag)

                cn = cmp.lower().replace(" ", "").replace("-", "")
                if "benzene" in cn and "ethyl" not in cn:
                    ev.benzene      = disp
                elif "toluene" in cn:
                    ev.toluene      = disp
                elif "ethylbenzene" in cn or "ethylb" in cn:
                    ev.ethylbenzene = disp
                elif "xylene" in cn:
                    ev.xylene       = disp
                elif "mtbe" in cn:
                    ev.mtbe         = disp
                elif "naphthalene" in cn:
                    ev.naphthalene  = disp
                elif "ph" == cn:
                    ev.ph           = disp
                elif "ec" in cn or "conductivity" in cn or "מוליכות" in cmp:
                    ev.ec           = disp
                elif "temp" in cn or "טמפרטורה" in cmp:
                    ev.temp         = disp
                elif "do" == cn or "oxygen" in cn or "חמצן" in cmp:
                    ev.do           = disp
                elif "turb" in cn or "עכירות" in cmp:
                    ev.turbidity    = disp
                elif "redox" in cn or "orp" in cn or "רדוקס" in cmp:
                    ev.redox        = disp
                elif "waterlevel" in cn or "עומקפניהמים" in cn.replace(" ",""):
                    ev.water_level  = disp
                elif "totaldepth" in cn or "עומקכללי" in cn.replace(" ",""):
                    ev.total_depth  = disp

            events.extend(grouped.values())

        except Exception as e:
            import traceback
            print(f"[annual_gw_report] parse_lab_reports error for {fname}: {e}")
            traceback.print_exc()

    return events


# ── Merge ─────────────────────────────────────────────────────────────────────

def merge_events(base: GWReportData, new_events: list[SamplingEvent]) -> GWReportData:
    """
    Add new lab events to base data, avoiding duplicates.
    """
    for ev in new_events:
        existing = base.get_event(ev.borehole, ev.date)
        if existing is None:
            base.events.append(ev)
            if ev.borehole not in base.boreholes:
                base.boreholes.append(ev.borehole)
        else:
            # Fill in any missing fields
            for attr in ["benzene", "toluene", "ethylbenzene", "xylene",
                         "mtbe", "naphthalene", "ph", "ec", "temp", "do",
                         "turbidity", "redox", "water_level", "total_depth",
                         "sample_depth"]:
                if getattr(existing, attr) is None and getattr(ev, attr) is not None:
                    setattr(existing, attr, getattr(ev, attr))
    return base


# ── Word report generator ─────────────────────────────────────────────────────

def generate_annual_report(
    data:          GWReportData,
    year:          str,
    author:        str       = "",
    approver:      str       = "",
    cutoff_date:   str       = "",   # last date of previous annual report
    template_bytes: bytes | None = None,
) -> bytes:
    """
    Generate annual Word report from GWReportData.
    Returns .docx bytes.
    """
    try:
        import docx as _docx
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        import subprocess, sys
        subprocess.check_call([sys.executable, "-m", "pip", "install",
                               "python-docx", "--break-system-packages", "-q"])
        import docx as _docx
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

    if template_bytes:
        doc = _docx.Document(io.BytesIO(template_bytes))
        _update_template(doc, data, year, author, approver, cutoff_date)
    else:
        doc = _build_from_scratch(data, year, author, approver, cutoff_date)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _update_template(doc, data, year, author, approver, cutoff_date):
    """Update an existing template doc in-place."""
    import docx as _docx
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    # Update cover page text
    for para in doc.paragraphs:
        if re.search(r'תקופתי|שנתי', para.text) and len(para.text) < 60:
            _replace_para_text(para, para.text.replace("תקופתי", "שנתי").replace(
                re.search(r'20\d\d', para.text).group() if re.search(r'20\d\d', para.text) else "", year))
        if re.search(r'20\d\d', para.text) and len(para.text) < 20:
            _replace_para_text(para, year)
        if "אוגוסט\|ינואר\|פברואר\|מרץ\|אפריל\|מאי\|יוני\|יולי\|ספטמבר\|אוקטובר\|נובמבר\|דצמבר" and \
           re.search(r'ינואר|פברואר|מרץ|אפריל|מאי|יוני|יולי|אוגוסט|ספטמבר|אוקטובר|נובמבר|דצמבר', para.text) \
           and len(para.text) < 30:
            _replace_para_text(para, f"דצמבר {year}")

    # Update author/approver in cover table
    for table in doc.tables[:3]:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if "מחבר" in " ".join(cells) and author:
                for c in row.cells:
                    if c.text.strip() and "מחבר" not in c.text and "חתימה" not in c.text:
                        _replace_cell_text(c, author)
            if ("מאשר" in " ".join(cells) or "בודק" in " ".join(cells)) and approver:
                for c in row.cells:
                    if c.text.strip() and "מאשר" not in c.text and "בודק" not in c.text and "חתימה" not in c.text:
                        _replace_cell_text(c, approver)

    # Replace all existing data tables
    tables_to_rebuild = []
    for ti, table in enumerate(doc.tables):
        header = " ".join(c.text for row in table.rows[:2] for c in row.cells)
        if any(k in header for k in ["בנזן", "MTBE", "Benzene"]):
            tables_to_rebuild.append(("lab", ti, table))
        elif any(k in header for k in ["pH", "עכירות", "רדוקס"]):
            tables_to_rebuild.append(("field", ti, table))

    for ttype, ti, table in tables_to_rebuild:
        # Clear existing data rows (keep header rows)
        header_rows = _count_header_rows(table)
        while len(table.rows) > header_rows:
            tr = table.rows[-1]._tr
            table._tbl.remove(tr)

        # Add new data rows
        boreholes = data.boreholes or sorted(set(e.borehole for e in data.events))
        if ttype == "lab":
            _fill_lab_table(table, data, boreholes, header_rows)
        else:
            _fill_field_table(table, data, boreholes, header_rows)

    # Update summary
    _update_summary(doc, data, year, cutoff_date)


def _build_from_scratch(data, year, author, approver, cutoff_date):
    """Build a complete annual report document from scratch."""
    import docx as _docx
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = _docx.Document()

    # Page setup (A4, RTL)
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = section.right_margin = Cm(2.5)
    section.top_margin  = section.bottom_margin = Cm(2.5)

    # Cover
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(f'דו"ח ניטור מי תהום שנתי')
    run.bold = True; run.font.size = Pt(18)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_r = sub_p.add_run(f"תחנת {data.site_name}" + (f", תחנה מספר {data.station_no}" if data.station_no else ""))
    sub_r.font.size = Pt(14)

    doc.add_paragraph(year).runs[0].font.size = Pt(14)
    doc.add_paragraph()

    # Author table
    auth_table = doc.add_table(rows=2, cols=2)
    auth_table.style = "Table Grid"
    auth_table.rows[0].cells[0].text = "חתימה"
    auth_table.rows[0].cells[1].text = "מחבר הדו\"ח"
    auth_table.rows[1].cells[1].text = author or data.author
    auth_table.rows[0].cells[0].text  # approver row
    appr_table = doc.add_table(rows=2, cols=2)
    appr_table.style = "Table Grid"
    appr_table.rows[0].cells[0].text = "חתימה"
    appr_table.rows[0].cells[1].text = "הדו\"ח מאשר"
    appr_table.rows[1].cells[1].text = approver or data.approver

    doc.add_page_break()

    # Title
    doc.add_heading(f'דו"ח ניטור מי תהום שנתי {year} – תחנת {data.site_name}', level=1)

    # Background (if available)
    if data.background_text:
        doc.add_heading("רקע", level=2)
        doc.add_paragraph(data.background_text)

    # Sampling description
    doc.add_heading("מהלך הדיגום ותוצאות בדיקות שדה ומעבדה", level=2)

    boreholes = data.boreholes or sorted(set(e.borehole for e in data.events))
    all_dates = data.all_dates_sorted()
    new_dates  = data.new_dates(cutoff_date) if cutoff_date else all_dates[-3:]

    # Find sampling date range for the year
    year_dates = [d for d in all_dates if year[-2:] in d or year in d]
    if year_dates:
        desc_date = year_dates[-1]  # most recent
        doc.add_paragraph(
            f"במהלך שנת {year} בוצעו דיגומי ניטור מי תהום בקידוחים "
            f"{', '.join(boreholes)}. הדיגום האחרון בוצע בתאריך {desc_date}."
        )

    # Table 1: Field measurements
    doc.add_heading("טבלה מספר 1: ממצאי בדיקות שדה", level=3)
    field_headers = ["שם קידוח", "תאריך", "pH", "EC\n[μS/cm]", "טמפרטורה\n[°C]",
                     "חמצן מומס\n[mg/l]", "עכירות\n[NTU]", "רדוקס\n[mV]",
                     "עומק פני המים\n[מ']", "עומק כללי\n[מ']", "עומק דגימה\n[מ']"]
    col_widths = [Cm(2), Cm(2.2)] + [Cm(1.8)] * 9
    ftable = doc.add_table(rows=1, cols=len(field_headers))
    ftable.style = "Table Grid"
    for i, (h, w) in enumerate(zip(field_headers, col_widths)):
        c = ftable.rows[0].cells[i]
        c.text = h
        c.width = w
        c.paragraphs[0].runs[0].bold = True

    for bh in boreholes:
        first = True
        for ev in data.events_for_borehole(bh):
            row = ftable.add_row()
            row.cells[0].text = bh if first else ""
            first = False
            row.cells[1].text  = ev.date or ""
            row.cells[2].text  = ev.ph or "-"
            row.cells[3].text  = ev.ec or "-"
            row.cells[4].text  = ev.temp or "-"
            row.cells[5].text  = ev.do or "-"
            row.cells[6].text  = ev.turbidity or "-"
            row.cells[7].text  = ev.redox or "-"
            row.cells[8].text  = ev.water_level or "-"
            row.cells[9].text  = ev.total_depth or "-"
            row.cells[10].text = ev.sample_depth or "-"

    doc.add_paragraph()

    # Table 2: Lab results
    doc.add_heading("טבלה מספר 2: תוצאות מעבדה – BTEX ו-MTBE", level=3)
    lab_headers = ["שם קידוח", "תאריך", "מפלס מים\n[מ']",
                   "בנזן\n[מג/ל]", "טולואן\n[מג/ל]",
                   "אתיל בנזן\n[מג/ל]", "כסילן\n[מג/ל]", "MTBE\n[מג/ל]"]
    ltable = doc.add_table(rows=1, cols=len(lab_headers))
    ltable.style = "Table Grid"
    for i, h in enumerate(lab_headers):
        c = ltable.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True

    # Target values row
    tv_row = ltable.add_row()
    tv_row.cells[0].text = "ערך יעד^"
    tv_row.cells[1].text = "-"
    tv_row.cells[2].text = "-"
    tv_row.cells[3].text = data.target_values.get("Benzene", "0.005")
    tv_row.cells[4].text = data.target_values.get("Toluene", "0.70")
    tv_row.cells[5].text = data.target_values.get("Ethyl Benzene", "0.30")
    tv_row.cells[6].text = data.target_values.get("Xylene", "0.50")
    tv_row.cells[7].text = data.target_values.get("MTBE", "0.04")

    for bh in boreholes:
        first = True
        for ev in data.events_for_borehole(bh):
            row = ltable.add_row()
            row.cells[0].text = bh if first else ""
            first = False
            row.cells[1].text = ev.date or ""
            row.cells[2].text = ev.water_level or "-"
            row.cells[3].text = ev.benzene      or "-"
            row.cells[4].text = ev.toluene      or "-"
            row.cells[5].text = ev.ethylbenzene or "-"
            row.cells[6].text = ev.xylene       or "-"
            row.cells[7].text = ev.mtbe         or "-"

    doc.add_paragraph("^ לפי תקן מי שתייה, 2013")
    doc.add_paragraph("- לא נבדק")

    # Summary
    doc.add_heading("סיכום והמלצות", level=2)
    _add_summary(doc, data, year, new_dates, boreholes)

    return doc


def _add_summary(doc, data, year, new_dates, boreholes):
    """Add auto-generated summary bullet points."""
    # Latest event per borehole
    latest_events = {}
    for bh in boreholes:
        evs = data.events_for_borehole(bh)
        if evs:
            latest_events[bh] = evs[-1]

    latest_date = max((e.date_obj() for e in latest_events.values() if e.date_obj()),
                      default=None)
    latest_date_str = latest_date.strftime("%d.%m.%y") if latest_date else "—"

    doc.add_heading("סיכום", level=3)
    bullets = [
        f"הניטור בשנת {year} כלל דיגומים בתאריכים: {', '.join(new_dates or ['—'])}.",
        f"הדיגום האחרון בוצע בתאריך {latest_date_str}.",
    ]
    # Findings per borehole
    for bh, ev in latest_events.items():
        finds = []
        for param, val in [("בנזן", ev.benzene), ("טולואן", ev.toluene),
                            ("אתיל בנזן", ev.ethylbenzene), ("כסילן", ev.xylene),
                            ("MTBE", ev.mtbe)]:
            if val and val not in ("-", "לא נבדק", "Not Detected", "ND"):
                finds.append(f"{param}: {val} מג/ל")
        if finds:
            bullets.append(f"בקידוח {bh} (דיגום {ev.date}): {', '.join(finds)}.")
        else:
            bullets.append(f"בקידוח {bh}: לא אותרו מרכיבי דלק מעל גבול הכימות.")

    for b in bullets:
        p = doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("המלצות", level=3)
    doc.add_paragraph("• מומלץ להמשיך לנטר את מי התהום בתחנה בהתאם להנחיות רשות המים (כל 120 יום).",
                       style="List Bullet")


def _update_summary(doc, data, year, cutoff_date):
    """Update the summary section in an existing template doc."""
    new_dates = data.new_dates(cutoff_date) if cutoff_date else []
    boreholes  = data.boreholes or sorted(set(e.borehole for e in data.events))
    latest_events = {}
    for bh in boreholes:
        evs = data.events_for_borehole(bh)
        if evs:
            latest_events[bh] = evs[-1]
    latest_date = max((e.date_obj() for e in latest_events.values() if e.date_obj()), default=None)
    latest_date_str = latest_date.strftime("%d.%m.%y") if latest_date else "—"

    # Find and update summary paragraphs
    in_summary = False
    for para in doc.paragraphs:
        if "סיכום" in para.text and ("המלצ" in para.text or len(para.text) < 20):
            in_summary = True
        if in_summary and re.search(r'הדיגום האחרון|הניטור ב', para.text):
            if "הדיגום האחרון" in para.text:
                _replace_para_text(para, f"הדיגום האחרון בוצע בתאריך {latest_date_str}.")
            if "הניטור ב" in para.text:
                _replace_para_text(para, f"הניטור בשנת {year} כלל דיגומים בתאריכים: {', '.join(new_dates or ['—'])}.")


# ── Helpers ───────────────────────────────────────────────────────────────────

def _normalize_date(s: str) -> str:
    s = s.replace("/", ".")
    parts = s.split(".")
    if len(parts) == 3:
        d, m, y = parts
        if len(y) == 4:
            y = y[2:]
        return f"{d.zfill(2)}.{m.zfill(2)}.{y}"
    return s


def _extract_date_from_sid_or_rec(rec: dict) -> str:
    """Extract date string from a lab record dict."""
    sampling = rec.get('sampling_date') or rec.get('date') or ''
    if sampling:
        import re as _re2
        sm = _re2.search(r'\d{1,2}[./]\d{1,2}[./]\d{2,4}', str(sampling))
        if sm:
            return _normalize_date(sm.group())
    sid = rec.get('sample_id', '')
    m = re.search(r'\d{1,2}[./]\d{1,2}[./]\d{2,4}', sid)
    if m:
        return _normalize_date(m.group())
    return sid



def _format_val(val, flag: str) -> str:
    if val is None:
        return "-"
    if flag in ("<LOQ", "<LOD", "<"):
        return f"<{val}"
    if str(val).upper() in ("NOT DETECTED", "ND", "N.D."):
        return f"<{val}"
    try:
        f = float(val)
        if f == int(f):
            return str(int(f))
        return str(round(f, 3))
    except (TypeError, ValueError):
        return str(val)


def _build_sample_borehole_map(fdata: bytes, fname: str) -> dict[str, str]:
    """
    Build a dict mapping sample_id → borehole_name from a Bactochem PDF.
    Uses regex over the full page text to find description+number pairs.
    """
    mapping: dict[str, str] = {}
    if not fname.lower().endswith(".pdf"):
        return mapping
    try:
        import fitz
        doc = fitz.open(stream=fdata, filetype="pdf")
        full_text = ""
        for page in doc:
            full_text += (page.get_text() or "") + "\n"

        # Find all description lines and their associated sample numbers
        # Pattern 1: desc line has borehole info, next non-empty line or nearby has מספר
        lines = full_text.split("\n")
        pending_bh = ""
        for i, line in enumerate(lines):
            stripped = line.strip()
            if not stripped:
                continue
            # Description line: contains תיאור/דיגום/בלנק AND borehole reference
            is_desc = any(k in stripped for k in ["תיאור", "דיגום", "בלנק"])
            bh_m = re.search(r'(?:הדוגמה:|הדוגמה)\s*:?\s*(\d+)-', stripped) or \
                   re.search(r'(\d+)-מת', stripped)
            if is_desc and bh_m:
                num = bh_m.group(1)
                pending_bh = f"מת-{num}"
                # Look ahead in next 5 lines for sample number
                for j in range(i + 1, min(i + 6, len(lines))):
                    num_m = re.search(r'מספר הדוגמה[:\s]+(\d+)', lines[j])
                    if num_m:
                        mapping[num_m.group(1).strip()] = pending_bh
                        pending_bh = ""
                        break
                # Also look backwards (RTL PDF may have swapped order)
                if pending_bh:
                    for j in range(max(0, i - 3), i):
                        num_m = re.search(r'מספר הדוגמה[:\s]+(\d+)', lines[j])
                        if num_m and num_m.group(1) not in mapping:
                            mapping[num_m.group(1).strip()] = pending_bh
                            pending_bh = ""
                            break
    except Exception:
        pass
    return mapping


def _extract_sampling_date_from_pdf(fdata: bytes) -> str:
    """Extract the sampling date from a Bactochem PDF."""
    try:
        import fitz
        doc = fitz.open(stream=fdata, filetype="pdf")
        for page in doc:
            text = page.get_text() or ""
            # "מועד דיגום: 04/12/2025 10:50" — date may be on same or next line
            # Also handle ":מועד דיגום" RTL
            for pattern in [
                r'מועד דיגום[:\s]+(\d{1,2}[./]\d{1,2}[./]\d{2,4})',
                r'(\d{1,2}[./]\d{1,2}[./]\d{4})\s+\d{1,2}:\d{2}',  # date+time on same line
            ]:
                m = re.search(pattern, text)
                if m:
                    return _normalize_date(m.group(1))
    except Exception:
        pass
    return ""
    # First try sampling_date field
    sampling = rec.get("sampling_date") or rec.get("date") or ""
    if sampling:
        sm = re.search(r'\d{1,2}[./]\d{1,2}[./]\d{2,4}', str(sampling))
        if sm:
            return _normalize_date(sm.group())
    # Try sample_id
    sid = rec.get("sample_id", "")
    m   = re.search(r'\d{1,2}[./]\d{1,2}[./]\d{2,4}', sid)
    if m:
        return _normalize_date(m.group())
    return sid


def _extract_borehole_from_rec(rec: dict) -> str:
    """
    Extract borehole name from a lab record.
    Handles:
      - sample_id like "מת-2" or "MW-3" directly
      - description like "דיגום מי תהום-סונול אור יהודה מת5-"
      - borehole field if present
    """
    # Direct borehole field
    bh = rec.get("borehole") or ""
    if bh and re.search(r'מת|MW|מד|B-\d|P-\d', bh, re.I):
        return bh.strip()

    # sample_id that looks like a borehole name
    sid = rec.get("sample_id", "")
    if re.search(r'^(מת|MW|מד|B-\d|P-\d)', sid, re.I):
        return sid.strip()

    # description / sample_description field (Bactochem PDF)
    desc = rec.get("description") or rec.get("sample_description") or \
           rec.get("sample_name") or ""
    if desc:
        # "דיגום מי תהום-סונול אור יהודה מת5-" → "מת-5"
        m = re.search(r'מת[\s\-]?(\d+)', desc)
        if m:
            return f"מת-{m.group(1)}"
        # "MW-3" / "MW3"
        m = re.search(r'\b(MW[\s\-]?\d+)\b', desc, re.I)
        if m:
            return m.group(1).upper().replace(" ", "-").replace("MW", "MW-").replace("MW--", "MW-")

    # Fallback: return sample_id as-is
    return sid


def _count_header_rows(table) -> int:
    """Estimate number of header rows in a table."""
    count = 0
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        if any(k in " ".join(cells) for k in ["pH", "יחידות", "בנזן", "שם קידוח", "Benzene", "MTBE"]):
            count += 1
        else:
            break
    return max(count, 1)


def _fill_lab_table(table, data: GWReportData, boreholes: list, header_rows: int):
    """Fill lab results into existing table."""
    import docx as _docx
    for bh in boreholes:
        first = True
        for ev in data.events_for_borehole(bh):
            row = table.add_row()
            row.cells[0].text = bh if first else ""
            first = False
            cells = row.cells
            # Try to match columns from header
            n = len(cells)
            if n >= 8:
                cells[1].text = ev.date or ""
                cells[2].text = ev.water_level or "-"
                cells[3].text = ev.benzene or "-"
                cells[4].text = ev.toluene or "-"
                cells[5].text = ev.ethylbenzene or "-"
                cells[6].text = ev.xylene or "-"
                cells[7].text = ev.mtbe or "-"
            elif n >= 6:
                cells[1].text = ev.date or ""
                cells[2].text = ev.benzene or "-"
                cells[3].text = ev.toluene or "-"
                cells[4].text = ev.xylene or "-"
                cells[5].text = ev.mtbe or "-"


def _fill_field_table(table, data: GWReportData, boreholes: list, header_rows: int):
    """Fill field measurements into existing table."""
    for bh in boreholes:
        first = True
        for ev in data.events_for_borehole(bh):
            row = table.add_row()
            cells = row.cells
            n = len(cells)
            cells[0].text = bh if first else ""
            first = False
            if n >= 10:
                cells[1].text  = ev.date or ""
                cells[2].text  = ev.ph or "-"
                cells[3].text  = ev.ec or "-"
                cells[4].text  = ev.temp or "-"
                cells[5].text  = ev.do or "-"
                cells[6].text  = ev.turbidity or "-"
                cells[7].text  = ev.redox or "-"
                cells[8].text  = ev.water_level or "-"
                cells[9].text  = ev.total_depth or "-"
                if n >= 11:
                    cells[10].text = ev.sample_depth or "-"


def _replace_para_text(para, new_text: str):
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)


def _replace_cell_text(cell, new_text: str):
    for para in cell.paragraphs:
        _replace_para_text(para, new_text)
        return


